import streamlit as st
import requests
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
import PyPDF2
import google.generativeai as genai
import openai
import os
import tempfile
import re
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import numpy as np
import hashlib
from typing import Dict, List
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import time
from requests.adapters import HTTPAdapter
from urllib3.util import Retry

# Connection pool and retry configuration for law.go.kr API
http_session = requests.Session()
http_session.headers.update({
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
    'Accept-Language': 'ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7',
})

retry_strategy = Retry(
    total=5,
    backoff_factor=1.5,
    status_forcelist=[429, 500, 502, 503, 504],
    raise_on_status=False
)
http_session.mount("http://", HTTPAdapter(max_retries=retry_strategy))
http_session.mount("https://", HTTPAdapter(max_retries=retry_strategy))

def request_get_with_retry(url, params=None, timeout=60):
    """API 요청시 ConnectionResetError 등을 방지하기 위한 재시도 및 지연(throttling) 헬퍼"""
    # 요청 간 0.15초 대기하여 API 서버에 무리를 주지 않고 차단 방지
    time.sleep(0.15)
    for attempt in range(3):
        try:
            response = http_session.get(url, params=params, timeout=timeout)
            response.raise_for_status()
            return response
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout, requests.exceptions.HTTPError, ConnectionResetError) as e:
            if attempt == 2:
                raise e
            # 점진적으로 대기시간을 늘리며 재시도 (1.5초, 3.0초)
            time.sleep((attempt + 1) * 1.5)

# Gemini File Search 통합
from gemini_file_search import (
    GeminiFileSearchManager,
    search_relevant_guidelines_gemini,
    search_violation_cases_gemini,
    get_gemini_store_manager
)


def _env_or_secret(key: str, default: str = "") -> str:
    """secrets.toml 이 없는 배포(Cloudtype 등)에서는 환경 변수만 사용."""
    v = os.environ.get(key)
    if v is not None and v != "":
        return v
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default


def _email_config_from_secrets() -> Dict:
    try:
        raw = st.secrets.get("email", {})
        return dict(raw) if raw else {}
    except Exception:
        return {}


st.set_page_config(
    layout="wide",
    initial_sidebar_state="expanded"
)

# Streamlit 기본 UI 숨기기
hide_streamlit_style = """
    <style>
    header {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# 페이지 설정
st.set_page_config(
    page_title="광역지자체 조례 검색, 비교, 분석",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 사용자 정의 CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #4f46e5, #7c3aed);
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        color: white;
        margin-bottom: 2rem;
    }
    .step-card {
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .result-card {
        background: #ffffff;
        border: 1px solid #d1d5db;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .law-title {
        color: #dc2626;
        font-weight: bold;
    }

    /* 탭 글자 크기 키우기 */
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
        font-size: 18px !important;
        font-weight: 600 !important;
        font-size: 1.1em;
        margin-bottom: 0.5rem;
    }
    .metro-name {
        color: #1e40af;
        font-weight: 600;
        margin-bottom: 0.3rem;
    }
    .stButton > button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

# API 설정
OC = "climsneys85"
search_url = "http://www.law.go.kr/DRF/lawSearch.do"
detail_url = "http://www.law.go.kr/DRF/lawService.do"

# 광역지자체 코드 및 이름
metropolitan_govs = {
    '6110000': '서울특별시',
    '6260000': '부산광역시',
    '6270000': '대구광역시',
    '6280000': '인천광역시',
    '6290000': '광주광역시',
    '6300000': '대전광역시',
    '5690000': '세종특별자치시',
    '6310000': '울산광역시',
    '6410000': '경기도',
    '6530000': '강원특별자치도',
    '6430000': '충청북도',
    '6440000': '충청남도',
    '6540000': '전북특별자치도',
    '6460000': '전라남도',
    '6470000': '경상북도',
    '6480000': '경상남도',
    '6500000': '제주특별자치도'
}

# 세션 상태 초기화
if 'search_results' not in st.session_state:
    st.session_state.search_results = []
if 'uploaded_pdf' not in st.session_state:
    st.session_state.uploaded_pdf = None
if 'search_query' not in st.session_state:
    st.session_state.search_query = ""
if 'word_doc_ready' not in st.session_state:
    st.session_state.word_doc_ready = False
if 'word_doc_data' not in st.session_state:
    st.session_state.word_doc_data = None
if 'selected_ordinances' not in st.session_state:
    st.session_state.selected_ordinances = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None

# Gemini File Search 관련 session state
if 'use_gemini_search' not in st.session_state:
    st.session_state.use_gemini_search = True  # 기본값: Gemini File Search 사용
if 'gemini_store_manager' not in st.session_state:
    st.session_state.gemini_store_manager = None

# Ollama Cloud 관련 session state
if 'use_ollama_cloud' not in st.session_state:
    st.session_state.use_ollama_cloud = True  # 기본값: Ollama Cloud 사용 (무료)
if 'ollama_api_key' not in st.session_state:
    st.session_state.ollama_api_key = _env_or_secret("OLLAMA_API_KEY", "")

# Cloudtype 등 환경 변수로 주입 (고급 설정 입력란과 동일 키 사용)
if "gemini_api_key_input" not in st.session_state:
    st.session_state.gemini_api_key_input = (
        _env_or_secret("GEMINI_API_KEY", "") or _env_or_secret("GOOGLE_API_KEY", "")
    )
if "openai_api_key_input" not in st.session_state:
    st.session_state.openai_api_key_input = _env_or_secret("OPENAI_API_KEY", "")

# RAG 벡터스토어 관련 session state
if 'rag_vectorstores' not in st.session_state:
    st.session_state.rag_vectorstores = None
if 'rag_loaded' not in st.session_state:
    st.session_state.rag_loaded = False

def load_rag_vectorstores():
    """PKL 파일에서 RAG 벡터스토어 로드"""
    import pickle

    if st.session_state.rag_loaded:
        return st.session_state.rag_vectorstores

    vectorstores = {}

    # 자치법규 매뉴얼 벡터스토어
    manual_path = "enhanced_vectorstore_20250914_101739.pkl"
    if os.path.exists(manual_path):
        try:
            with open(manual_path, 'rb') as f:
                vectorstores['manual'] = pickle.load(f)
            st.success(f"✅ 자치법규 매뉴얼 로드 완료")
        except Exception as e:
            st.warning(f"⚠️ 자치법규 매뉴얼 로드 실패: {e}")

    # 재의·제소 조례 모음집 벡터스토어
    cases_path = "3. 지방자치단체의 재의·제소 조례 모음집(Ⅸ) (1)_new_vectorstore.pkl"
    if os.path.exists(cases_path):
        try:
            with open(cases_path, 'rb') as f:
                vectorstores['cases'] = pickle.load(f)
            st.success(f"✅ 재의·제소 판례 모음집 로드 완료")
        except Exception as e:
            st.warning(f"⚠️ 재의·제소 판례 모음집 로드 실패: {e}")

    st.session_state.rag_vectorstores = vectorstores
    st.session_state.rag_loaded = True
    return vectorstores

def search_rag_context(query, vectorstores, top_k=5):
    """RAG 벡터스토어에서 관련 문서 검색"""
    results = []

    # 품질 필터 함수: 목차/제목만 있는 청크 제외
    def is_quality_content(text):
        """유용한 내용인지 판단"""
        # 최소 길이 체크 (100자 미만은 목차일 가능성 높음)
        if len(text) < 100:
            return False

        # 목차/제목 패턴 감지
        toc_patterns = [
            r'^제\d+장\s+',  # 제1장
            r'^제\d+절\s+',  # 제1절
            r'^\d+\.\s+\w+\s*$',  # 1. 제목
            r'^[가-힣]+\s+\d+$',  # 목차 번호
            r'^\s*목\s*차\s*$',  # 목차
            r'^\s*차\s*례\s*$',  # 차례
        ]

        for pattern in toc_patterns:
            if re.search(pattern, text.strip(), re.MULTILINE):
                # 패턴이 있어도 내용이 충분히 있으면 허용
                if len(text) > 300:
                    return True
                return False

        # 문장 완성도 체크: 마침표가 3개 이상 있어야 함 (설명이 있는 텍스트)
        sentence_count = text.count('.') + text.count('다.') + text.count('함.')
        if sentence_count < 2:
            return False

        # 실제 법률 용어나 설명이 포함되어 있는지
        useful_keywords = ['판단', '해석', '따라서', '경우', '규정', '위반', '적법', '위법', '검토', '사례', '판례']
        has_useful_content = any(kw in text for kw in useful_keywords)

        return has_useful_content or len(text) > 500

    for store_name, store_data in vectorstores.items():
        try:
            # 벡터스토어 형식에 따라 검색 수행
            if isinstance(store_data, dict):
                # chunks 키가 있는 경우 (우선 사용)
                if 'chunks' in store_data:
                    chunks = store_data['chunks']
                    query_keywords = [kw.lower() for kw in query.split() if len(kw) > 1]

                    scored_chunks = []
                    for chunk in chunks:
                        if isinstance(chunk, dict) and 'text' in chunk:
                            text = chunk['text']
                        elif isinstance(chunk, str):
                            text = chunk
                        else:
                            continue

                        # 품질 필터: 유용한 내용인지 체크
                        if not is_quality_content(text):
                            continue

                        # 키워드 매칭 점수 계산
                        text_lower = text.lower()
                        keyword_score = sum(1 for kw in query_keywords if kw in text_lower)

                        # 내용 밀도 보너스: 긴 텍스트에 보너스 점수
                        length_bonus = min(len(text) / 500, 3.0)  # 최대 3점 보너스

                        # 법률 분석 키워드 보너스
                        analysis_keywords = ['판단', '검토', '위법', '적법', '사례', '판례', '해석', '기준']
                        analysis_bonus = sum(0.5 for kw in analysis_keywords if kw in text)

                        total_score = keyword_score + length_bonus + analysis_bonus

                        if keyword_score > 0:
                            scored_chunks.append((text, total_score))

                    # 상위 결과 선택
                    scored_chunks.sort(key=lambda x: x[1], reverse=True)
                    for text, score in scored_chunks[:top_k]:
                        results.append({
                            'source': store_name,
                            'text': text[:2000],  # 최대 2000자
                            'score': score
                        })

                # texts 키가 있는 경우
                elif 'texts' in store_data:
                    texts = store_data['texts']
                    query_keywords = [kw.lower() for kw in query.split() if len(kw) > 1]

                    scored_texts = []
                    for text in texts:
                        if isinstance(text, str):
                            text_lower = text.lower()
                            score = sum(1 for kw in query_keywords if kw in text_lower)
                            if score > 0:
                                scored_texts.append((text, score))

                    scored_texts.sort(key=lambda x: x[1], reverse=True)
                    for text, score in scored_texts[:top_k]:
                        results.append({
                            'source': store_name,
                            'text': text[:2000],
                            'score': score
                        })

                # documents 키가 있는 경우
                elif 'documents' in store_data:
                    docs = store_data['documents']
                    query_keywords = [kw.lower() for kw in query.split() if len(kw) > 1]

                    scored_docs = []
                    for doc in docs:
                        if isinstance(doc, dict):
                            text = doc.get('text', doc.get('content', ''))
                        elif isinstance(doc, str):
                            text = doc
                        else:
                            continue

                        text_lower = text.lower()
                        score = sum(1 for kw in query_keywords if kw in text_lower)
                        if score > 0:
                            scored_docs.append((text, score))

                    scored_docs.sort(key=lambda x: x[1], reverse=True)
                    for text, score in scored_docs[:top_k]:
                        results.append({
                            'source': store_name,
                            'text': text[:2000],
                            'score': score
                        })
            elif hasattr(store_data, 'similarity_search'):
                # LangChain 스타일 벡터스토어
                docs = store_data.similarity_search(query, k=top_k)
                for doc in docs:
                    results.append({
                        'source': store_name,
                        'text': doc.page_content[:2000],
                        'score': 1.0
                    })
        except Exception as e:
            st.warning(f"⚠️ {store_name} 검색 중 오류: {e}")

    # 점수순 정렬
    results.sort(key=lambda x: x.get('score', 0), reverse=True)
    return results[:top_k * 2]  # 최대 top_k * 2개 반환

def call_ollama_cloud_api(prompt, model="gpt-oss:120b-cloud", max_chars=100000):
    """Ollama Cloud API를 호출하여 텍스트 생성

    Args:
        prompt: 분석 프롬프트
        model: 사용할 모델 (기본: gpt-oss:120b-cloud)
        max_chars: 최대 문자 수 (기본: 100000자, 한글 기준 약 50-70K 토큰)
    """
    try:
        api_key = st.session_state.ollama_api_key
        if not api_key or api_key == "YOUR_OLLAMA_API_KEY_HERE":
            st.error("Ollama Cloud API 키가 설정되지 않았습니다.")
            return None

        # 프롬프트 길이 제한 (토큰 제한 방지 - 한글은 토큰 효율이 낮음)
        original_len = len(prompt)
        if original_len > max_chars:
            st.warning(f"⚠️ 프롬프트가 너무 깁니다 ({original_len:,}자). {max_chars:,}자로 자동 축소합니다.")

            # 섹션 마커를 찾아서 지능형 축소
            # 1. 법리적 가이드라인 (최우선 보존)
            # 2. 검토 대상 조례 원문 (필수 보존)
            # 3. 상위법령 (부분 축소 가능)
            # 4. RAG 참고자료 (부분 축소 가능)
            # 5. 분석 지시사항 (필수 보존)

            try:
                # 주요 섹션 경계 찾기
                guideline_start = prompt.find("🚨 **필독: 조례 위법 판단")
                ordinance_start = prompt.find("📄 **[검토 대상 조례 원문 시작]**")
                ordinance_end = prompt.find("📄 **[검토 대상 조례 원문 종료]**")
                reference_start = prompt.find("📚 **[참고자료:")
                reference_end = prompt.find("📚 **[참고자료 종료]**")
                analysis_instruction_start = prompt.find("아래 기준에 따라 분석해줘")

                # 필수 섹션 추출
                guideline_section = prompt[guideline_start:ordinance_start] if guideline_start != -1 and ordinance_start != -1 else ""
                ordinance_section = prompt[ordinance_start:ordinance_end + 100] if ordinance_start != -1 and ordinance_end != -1 else ""

                # 조례 원문이 너무 길면 일부 축소 (앞부분 유지)
                if len(ordinance_section) > max_chars * 0.4:
                    ordinance_header = ordinance_section[:2000]  # 헤더 보존
                    ordinance_content_limit = int(max_chars * 0.4) - 2000
                    ordinance_section = ordinance_header + ordinance_section[2000:2000+ordinance_content_limit] + "\n\n... [조례 일부 생략] ...\n\n" + ordinance_section[-500:]

                # 참고자료는 요약 (첫 5개 항목만)
                reference_section = ""
                if reference_start != -1 and reference_end != -1:
                    ref_content = prompt[reference_start:reference_end + 100]
                    # 참고자료 개수 제한
                    ref_items = ref_content.split("[참고자료")
                    if len(ref_items) > 6:  # 헤더 + 5개 항목
                        reference_section = "[참고자료".join(ref_items[:6]) + "\n\n... [참고자료 일부 생략 - 위법 판단 근거로만 사용] ...\n\n📚 **[참고자료 종료]**\n" + "=" * 80 + "\n"
                    else:
                        reference_section = ref_content

                # 분석 지시사항 (필수 완전 보존)
                if analysis_instruction_start != -1:
                    instruction_section = prompt[analysis_instruction_start:]  # 분석 지시사항 전체 보존
                else:
                    # 찾지 못하면 마지막 20% 보존 (안전장치)
                    instruction_section = prompt[-int(max_chars * 0.2):]

                # 중간 섹션 (상위법령, 타시도 조례) - 남은 공간만큼 할당
                if ordinance_end != -1 and reference_start != -1:
                    middle_section = prompt[ordinance_end + 100:reference_start]
                    # 중간 섹션 크기 제한 (최대 20%)
                    max_middle = int(max_chars * 0.2)
                    if len(middle_section) > max_middle:
                        middle_section = middle_section[:max_middle] + "\n\n... [상위법령/타시도 조례 일부 생략] ...\n\n"
                else:
                    middle_section = ""

                # 재조립
                prompt = guideline_section + ordinance_section + middle_section + reference_section + instruction_section

                st.info(f"✅ 프롬프트를 {len(prompt):,}자로 축소했습니다 (필수 섹션 보존: 법리 가이드라인, 조례 원문, 분석 지시사항)")

            except Exception as e:
                # 섹션 파싱 실패 시 기존 방식 사용
                st.warning(f"⚠️ 지능형 축소 실패, 단순 축소 적용: {str(e)}")
                front_chars = int(max_chars * 0.5)  # 앞부분 50%
                back_chars = int(max_chars * 0.3)   # 뒷부분 30%
                prompt = (
                    prompt[:front_chars] +
                    f"\n\n... [중략: 원본 {original_len:,}자 중 {original_len - max_chars:,}자 생략됨] ...\n\n" +
                    prompt[-back_chars:]
                )
                st.info(f"✅ 프롬프트를 {len(prompt):,}자로 축소했습니다.")

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "stream": False
        }

        response = requests.post(
            "https://ollama.com/api/chat",
            headers=headers,
            json=payload,
            timeout=180  # 타임아웃 증가 (긴 프롬프트 처리)
        )

        if response.status_code == 200:
            result = response.json()
            # Ollama API 응답 형식에서 텍스트 추출
            if "message" in result and "content" in result["message"]:
                return result["message"]["content"]
            elif "response" in result:
                return result["response"]
            else:
                st.warning(f"예상치 못한 응답 형식: {result}")
                return str(result)
        else:
            st.error(f"Ollama Cloud API 오류: {response.status_code} - {response.text}")
            return None

    except requests.exceptions.Timeout:
        st.error("Ollama Cloud API 요청 시간 초과 (120초)")
        return None
    except Exception as e:
        st.error(f"Ollama Cloud API 호출 오류: {str(e)}")
        return None

def get_ordinance_detail(ordinance_id):
    """조례 상세 내용 가져오기"""
    params = {
        'OC': OC,
        'target': 'ordin',
        'ID': ordinance_id,
        'type': 'XML'
    }
    try:
        response = request_get_with_retry(detail_url, params=params, timeout=60)
        root = ET.fromstring(response.text)
        articles = []
        for article in root.findall('.//조'):
            content = article.find('조내용').text if article.find('조내용') is not None else ""
            if content:
                content = content.replace('<![CDATA[', '').replace(']]>', '')
                content = content.replace('<p>', '').replace('</p>', '\n')
                content = content.replace('<br/>', '\n')
                content = content.replace('<br>', '\n')
                content = content.replace('&nbsp;', ' ')
                content = content.strip()
            if content:
                articles.append(content)
        return articles
    except Exception:
        return []

def search_ordinances(query):
    """조례 검색 함수"""
    results = []
    total_count = 0
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    total_metros = len(metropolitan_govs)
    
    for idx, (org_code, metro_name) in enumerate(metropolitan_govs.items()):
        status_text.text(f"검색 중... {metro_name} ({idx + 1}/{total_metros})")
        progress_bar.progress((idx + 1) / total_metros)
        
        try:
            params = {
                'OC': OC,
                'target': 'ordin',
                'type': 'XML',
                'query': query,
                'display': 100,
                'search': 1,
                'sort': 'ddes',
                'page': 1,
                'org': org_code
            }
            
            response = request_get_with_retry(search_url, params=params, timeout=60)
            response.raise_for_status()
            
            root = ET.fromstring(response.text)
            
            for law in root.findall('.//law'):
                ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                
                if 기관명 != metro_name:
                    continue
                
                # 검색어 매칭 로직
                search_terms = [term.lower() for term in query.split() if term.strip()]
                ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                if not all(term in ordinance_name_clean for term in search_terms):
                    continue
                
                total_count += 1
                articles = get_ordinance_detail(ordinance_id)
                
                results.append({
                    'name': ordinance_name,
                    'content': articles,
                    'metro': metro_name
                })
                
        except Exception as e:
            st.warning(f"검색 중 오류 발생 ({metro_name}): {str(e)}")
            continue
    
    progress_bar.empty()
    status_text.empty()
    
    return results, total_count

def create_word_document(query, results):
    """Word 문서 생성 함수"""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)

    # 제목 추가
    title = doc.add_heading('조례 검색 결과', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'검색어: {query}')
    doc.add_paragraph(f'총 {len(results)}건의 조례가 검색되었습니다.\n')

    # 조례를 3개씩 그룹화하여 3단 비교표 형태로 생성
    for i in range(0, len(results), 3):
        current_laws = results[i:i+3]
        while len(current_laws) < 3:
            current_laws.append({'name': '', 'content': [], 'metro': ''})

        # 표 생성 (1행, 3열 고정)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = True

        # 각 셀에 조례 내용 추가
        for idx, law in enumerate(current_laws):
            cell = table.cell(0, idx)
            paragraph = cell.paragraphs[0]
            
            if law['name']:
                # 조례명 추가 (지자체명 + 조례명)
                run = paragraph.add_run(f"{law['metro']}\n{law['name']}\n\n")
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                
                # 조문 내용 추가
                if law['content']:
                    content_text = '\n\n'.join(law['content'])
                    paragraph.add_run(content_text)
                else:
                    paragraph.add_run('(조문 없음)')

        # 마지막 페이지가 아니면 페이지 나누기 추가
        if i + 3 < len(results):
            doc.add_page_break()

    return doc

def extract_pdf_text(pdf_file):
    """PDF 텍스트 추출 함수"""
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
        return text
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류 발생: {str(e)}")
        return None

def extract_superior_laws(pdf_text):
    """조례안에서 상위법령 추출 함수 - GUI 검증된 로직 적용"""
    import re

    # 상위법 후보 추출을 위한 키워드 (조례안에서 상위법령 언급하는 모든 맥락 포함)
    law_check_keywords = [
        '위반', '위배', '충돌', '저촉', '준수', '적합', '불일치',
        '상위법', '상위 법령', '상위법령', '법령과의 관계', '법령과의 충돌', '법령과의 위배',
        '관계법령', '근거법령', '법적근거', '참고사항', '관련법령', '소관법령',
        '법령', '법률', '시행령', '시행규칙', '규정', '개정', '제정', '법'  # 일반적인 법령 언급
    ]

    # 법령명 패턴 (시행령/시행규칙 추출 개선)
    law_pattern = re.compile(r'([가-힣\w\s]*(?:법|시행령|시행규칙))\s*(?:[」]|$|[.,;:\s])', re.MULTILINE)

    # 상위법 후보 추출
    upper_law_candidates = set()

    # 1. 상위법 관련 맥락이 있는 줄에서 법령명 추출
    for line in pdf_text.split('\n'):
        if any(keyword in line for keyword in law_check_keywords):
            for match in law_pattern.finditer(line):
                law_name = match.group(1).strip()
                if law_name:
                    upper_law_candidates.add(law_name)

    # 2. 추가 패턴: 「법령명」 형식으로 따옴표 안에 있는 법령명 추출
    quote_pattern = re.compile(r'[「『]([^」』]*(?:법|시행령|시행규칙))[」』]')
    for match in quote_pattern.finditer(pdf_text):
        law_name = match.group(1).strip()
        if law_name:
            upper_law_candidates.add(law_name)

    # 3. 추가 패턴: "○○법령:" 또는 "관계법령:" 뒤에 오는 법령명
    relation_pattern = re.compile(r'(?:관계법령|근거법령|법적근거|소관법령|관련법령)\s*[:：]\s*[「『]?([^」』\n]*(?:법|시행령|시행규칙))[」』]?')
    for match in relation_pattern.finditer(pdf_text):
        law_name = match.group(1).strip()
        if law_name:
            upper_law_candidates.add(law_name)

    # 불용어 리스트 (실존하지 않는 법령명)
    invalid_law_names = {
        '자치입법', '조례', '규칙', '지침', '내규', '예규', '훈령', '적법',
        '입법', '상위법', '위법', '합법', '불법', '방법', '헌법상', '헌법적',
        '법적', '법률적', '법령상', '법률상', '법률', '법령', '법', '규정',
        '조항', '조문', '규범', '원칙', '기준', '사항', '내용', '관련법',
        '관련 법', '관련법령', '관련 법령'
    }

    def is_valid_law_name(name):
        """유효한 법령명인지 검증"""
        # 대소문자, 공백 모두 제거 후 비교
        name_clean = name.strip().replace(' ', '').lower()

        # 불용어 체크
        for invalid in invalid_law_names:
            if name_clean == invalid.replace(' ', '').lower():
                return False

        # 숫자+법(예: 1법, 2법 등)도 제외
        if name_clean and name_clean[0].isdigit():
            return False

        # 너무 짧은 이름 제외
        if len(name_clean) < 3:
            return False

        return True

    # 유효한 법령명만 필터링
    valid_laws = []
    for law_name in upper_law_candidates:
        if is_valid_law_name(law_name):
            valid_laws.append(law_name)

    # 🆕 시행령/시행규칙 자동 유추 추가
    additional_laws = []
    for law in valid_laws:
        if law.endswith('법') and '시행' not in law:
            # 해당 법률의 시행령과 시행규칙을 자동으로 추가
            base_name = law

            # 시행령 추가 (일반적인 패턴)
            potential_decree = f"{base_name} 시행령"
            if potential_decree not in valid_laws:
                additional_laws.append(potential_decree)

            # 시행규칙 추가 (일반적인 패턴)
            potential_rule = f"{base_name} 시행규칙"
            if potential_rule not in valid_laws:
                additional_laws.append(potential_rule)

    # 추가된 법령들을 포함
    if additional_laws:
        import streamlit as st
        st.info(f"🔄 자동 추가된 하위 법령: {len(additional_laws)}개")
        with st.expander("📋 자동 추가된 법령", expanded=False):
            for law in additional_laws:
                st.markdown(f"- {law}")
        valid_laws.extend(additional_laws)

    # 중복 제거 및 정렬
    unique_laws = list(set(valid_laws))
    unique_laws.sort()

    return unique_laws[:20]  # 최대 20개 반환

def get_superior_law_content_xml(law_name):
    """XML API를 통해 상위법령 내용 가져오기 (성공적인 로직 적용)"""
    try:
        import xml.etree.ElementTree as ET
        import re

        # 검색어 최적화: 띄어쓰기와 특수문자 정리
        search_query = law_name.strip()

        # 1단계: 법령 검색 (더 많은 결과 반환)
        search_params = {
            'OC': OC,
            'target': 'law',
            'type': 'XML',
            'query': search_query,
            'display': 10  # 더 많은 결과 검색
        }
        
        search_response = request_get_with_retry(search_url, params=search_params, timeout=30)
        if search_response.status_code != 200:
            return get_superior_law_content_xml_fallback(law_name)
        
        search_root = ET.fromstring(search_response.text)
        
        # 현행 법령 찾기 - 더 유연한 검색
        current_laws = []
        for law in search_root.findall('.//law'):
            status = law.find('현행연혁코드')
            if status is not None and status.text == '현행':
                law_id_elem = law.find('법령ID')
                law_name_elem = law.find('법령명한글')
                if law_id_elem is not None and law_name_elem is not None:
                    current_laws.append({
                        'id': law_id_elem.text,
                        'name': law_name_elem.text
                    })

        if not current_laws:
            return get_superior_law_content_xml_fallback(law_name)
        
        # 가장 관련성 높은 법령 선택 (개선된 매칭 알고리즘)
        best_law = None
        best_score = -1

        for law_info in current_laws:
            found_name = law_info['name']
            score = 0

            # 1. 정확한 매칭 우선
            if found_name == law_name:
                score += 1000

            # 2. 부분 매칭 점수 (양방향)
            if law_name in found_name:
                score += 500
            if found_name in law_name:
                score += 300

            # 3. 핵심 키워드 매칭 (개선된 로직)
            law_lower = law_name.lower().replace(' ', '')
            found_lower = found_name.lower().replace(' ', '')

            # 여객자동차 운수사업법 관련 특별 점수
            if '여객자동차' in law_lower and '운수사업' in law_lower:
                if '여객자동차' in found_lower and '운수사업' in found_lower:
                    score += 400  # 여객자동차 운수사업법 관련 높은 점수
                    if '시행규칙' in law_lower and '시행규칙' in found_lower:
                        score += 200  # 시행규칙 매칭 추가 점수

            # 도로교통법 관련
            if '도로' in law_lower and '교통' in law_lower:
                if '도로교통' in found_lower and '특별회계' not in found_lower:
                    score += 300
                elif '교통시설' in found_lower:
                    score -= 100

            # 4. 법령 유형 매칭 점수 (요청된 유형과 일치하는지)
            requested_type = ''
            if '시행규칙' in law_lower:
                requested_type = '시행규칙'
            elif '시행령' in law_lower:
                requested_type = '시행령'
            elif '법' in law_lower and '시행' not in law_lower:
                requested_type = '법'

            if requested_type:
                if requested_type in found_lower:
                    score += 300  # 요청된 법령 유형과 일치하면 높은 점수
                elif requested_type == '법' and found_lower.endswith('법') and '시행' not in found_lower:
                    score += 300
            else:
                # 기본 우선순위 (법률 > 시행령 > 시행규칙)
                if found_lower.endswith('법') and not ('시행령' in found_lower or '시행규칙' in found_lower):
                    score += 100
                elif '시행령' in found_lower:
                    score += 50
                elif '시행규칙' in found_lower:
                    score += 25

            # 5. 길이 페널티 완화 (너무 긴 법령명은 약간 감점)
            if len(found_name) > 30:
                score -= 30

            if score > best_score:
                best_score = score
                best_law = law_info
        
        if best_law:
            law_id = best_law['id']
            exact_law_name = best_law['name']
        else:
            # 폴백: 첫 번째 법령
            law_id = current_laws[0]['id']
            exact_law_name = current_laws[0]['name']

        if not law_id:
            return get_superior_law_content_xml_fallback(law_name)
        
        # 2단계: 상세 정보 가져오기
        detail_params = {
            'OC': OC,
            'target': 'law',
            'type': 'XML',
            'ID': law_id
        }
        
        detail_response = request_get_with_retry(detail_url, params=detail_params, timeout=30)
        if detail_response.status_code != 200:
            return get_superior_law_content_xml_fallback(law_name)

        detail_root = ET.fromstring(detail_response.text)
        
        # 3단계: 성공적인 추출 로직 적용 - 연결된 본문으로 처리
        upper_law_text = ""
        jo_count = 0
        hang_count = 0 
        ho_count = 0
        
        for node in detail_root.iter():
            if node.tag == '조문내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += content + '\n'
                jo_count += 1
            elif node.tag == '항내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += '    ' + content + '\n'
                hang_count += 1
            elif node.tag == '호내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += '        ' + content + '\n'
                ho_count += 1

        if upper_law_text.strip():
            # 스마트 필터링: 조례 관련 키워드가 포함된 부분 우선 추출
            def smart_filter_content(content, max_length=50000):
                """조례와 관련성 높은 부분을 우선 추출"""
                lines = content.split('\n')
                
                # 조례 관련 키워드 (도로교통법 관련)
                priority_keywords = [
                    '시장', '군수', '구청장', '지방자치단체', '조례', '시도', '시군구',
                    '위임', '위탁', '권한', '사무', '신고', '허가', '승인', '지정',
                    '주차', '정차', '금지', '제한', '구역', '시설', '설치', '관리'
                ]
                
                # 우선순위별로 라인 분류
                high_priority = []
                medium_priority = []
                low_priority = []
                
                for line in lines:
                    line_lower = line.lower()
                    priority_count = sum(1 for keyword in priority_keywords if keyword in line_lower)
                    
                    if priority_count >= 2:
                        high_priority.append(line)
                    elif priority_count >= 1:
                        medium_priority.append(line)
                    else:
                        low_priority.append(line)
                
                # 우선순위별로 결합
                filtered_content = []
                current_length = 0
                
                # 1단계: 높은 우선순위
                for line in high_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                # 2단계: 중간 우선순위
                for line in medium_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                # 3단계: 낮은 우선순위 (공간이 남으면)
                for line in low_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                result = '\n'.join(filtered_content)
                if len(content) > len(result):
                    result += "\n\n[... 조례 관련성이 높은 부분을 우선 표시하였습니다 ...]"
                
                return result
            
            # 스마트 필터링 적용 (Gemini 2.0 flash exp는 더 큰 컨텍스트 지원)
            max_length = 80000
            if len(upper_law_text) > max_length:
                truncated_text = smart_filter_content(upper_law_text, max_length)
            else:
                truncated_text = upper_law_text.strip()
            
            # 모든 조문을 하나의 연결된 본문으로 처리
            result = {
                'law_name': exact_law_name,
                'law_id': law_id,
                'content': truncated_text
            }

            return result
        else:
            return get_superior_law_content_xml_fallback(law_name)

    except Exception as e:
        return get_superior_law_content_xml_fallback(law_name)

def get_superior_law_content_xml_fallback(law_name):
    """XML 방식 폴백 (간소화 버전)"""
    try:
        search_params = {
            'OC': OC,
            'target': 'law',
            'type': 'XML',
            'query': law_name,
            'display': 5,
            'search': 1
        }

        search_response = request_get_with_retry(search_url, params=search_params, timeout=30)
        
        if search_response.status_code != 200:
            return None

        if not search_response.text.strip():
            return None

        try:
            search_root = ET.fromstring(search_response.text)
        except ET.ParseError as xml_err:
            return None
        
        law_id = None
        exact_law_name = None
        
        for law in search_root.findall('.//law'):
            found_name = law.find('법령명').text if law.find('법령명') is not None else ""
            found_id = law.find('법령ID').text if law.find('법령ID') is not None else None
            
            if found_name == law_name or (law_name in found_name):
                law_id = found_id
                exact_law_name = found_name
                break
        
        if not law_id:
            return None
        
        detail_params = {
            'OC': OC,
            'target': 'law', 
            'ID': law_id,
            'type': 'XML'
        }
        
        detail_response = request_get_with_retry(detail_url, params=detail_params, timeout=30)
        detail_root = ET.fromstring(detail_response.text)
        
        articles = []
        for article in detail_root.findall('.//조'):
            article_num = article.find('조문번호').text if article.find('조문번호') is not None else ""
            article_title = article.find('조문제목').text if article.find('조문제목') is not None else ""
            article_content = article.find('조문내용').text if article.find('조문내용') is not None else ""
            
            if article_content:
                article_content = article_content.replace('<![CDATA[', '').replace(']]>', '')
                article_content = article_content.replace('<p>', '').replace('</p>', '\n')
                article_content = article_content.replace('<br/>', '\n').replace('<br>', '\n')
                article_content = article_content.replace('&nbsp;', ' ')
                article_content = article_content.strip()
                
                if article_content:
                    articles.append({
                        'number': article_num,
                        'title': article_title,
                        'content': article_content
                    })
        
        return {
            'law_name': exact_law_name,
            'law_id': law_id,
            'articles': articles
        }
        
    except Exception as e:
        return None

# 기존 함수를 새 XML 방식으로 교체
def get_superior_law_content(law_name):
    """상위법령 내용 가져오기 (XML 방식)"""
    return get_superior_law_content_xml(law_name)

def normalize_law_name(law_name):
    """법령명을 정규화하여 중복 제거"""
    import re

    # 1. 기본 정리: 앞뒤 공백 제거
    normalized = law_name.strip()

    # 2. 과도한 띄어쓰기 제거 (2개 이상의 공백을 1개로)
    normalized = re.sub(r'\s+', ' ', normalized)

    # 3. 특정 패턴 정규화
    # "관광진흥 법" -> "관광진흥법"
    normalized = re.sub(r'(\w+)\s+(법|령|규칙)$', r'\1\2', normalized)

    # 4. 폐광지역개발지원 관련 법령 정규화
    if '폐광지' in normalized or '역개발' in normalized:
        if '특별법' in normalized:
            normalized = "폐광지역개발지원에관한특별법"

    # 5. 너무 짧은 법령명 제거 (3글자 이하)
    if len(normalized) <= 3:
        return None

    # 6. 명확히 잘못된 추출 제거
    invalid_patterns = [
        r'^한특별법$',  # "한특별법"
        r'^\w{1,2}특별법$',  # 너무 짧은 특별법
    ]

    for pattern in invalid_patterns:
        if re.match(pattern, normalized):
            return None

    return normalized

def group_laws_by_hierarchy(superior_laws):
    """법령을 계층별로 그룹화하는 함수 (정규화 적용)"""
    law_groups = {}

    # 1단계: 법령명 정규화 및 중복 제거
    normalized_laws = set()
    for law_name in superior_laws:
        normalized = normalize_law_name(law_name)
        if normalized:  # None이 아닌 경우만 추가
            normalized_laws.add(normalized)

    if len(superior_laws) != len(normalized_laws):
        import streamlit as st
        st.info(f"🔧 법령명 정규화: {len(superior_laws)}개 → {len(normalized_laws)}개로 중복 제거")

        # 제거된 중복 법령 표시
        removed_laws = []
        for original in superior_laws:
            normalized = normalize_law_name(original)
            if not normalized or (normalized != original and normalized in normalized_laws):
                removed_laws.append(original)

        if removed_laws:
            with st.expander("🗑️ 제거된 중복/잘못된 법령명", expanded=False):
                for removed in removed_laws:
                    st.markdown(f"- {removed}")

    # 정규화 과정 로깅
    for original in superior_laws:
        normalized = normalize_law_name(original)

    # 2단계: 정규화된 법령명으로 그룹화
    for law_name in normalized_laws:
        # 기본 법령명 추출 (시행령, 시행규칙 제거)
        base_name = law_name
        law_type = 'law'  # 기본값: 법률
        
        if '시행규칙' in law_name:
            base_name = law_name.replace(' 시행규칙', '').replace('시행규칙', '')
            law_type = 'rule'
        elif '시행령' in law_name:
            base_name = law_name.replace(' 시행령', '').replace('시행령', '')
            law_type = 'decree'
        elif law_name.endswith('령') and not law_name.endswith('법령'):
            law_type = 'decree'
        elif law_name.endswith('규칙'):
            law_type = 'rule'
            
        # 그룹에 추가
        if base_name not in law_groups:
            law_groups[base_name] = {'law': None, 'decree': None, 'rule': None}
        
        law_groups[base_name][law_type] = law_name
    
    return law_groups

def get_all_superior_laws_content(superior_laws):
    """모든 상위법령 내용을 가져오는 함수 - 계층별 그룹화"""
    superior_laws_content = []
    
    if not superior_laws:
        return superior_laws_content
    
    # 1단계: 법령을 계층별로 그룹화
    law_groups = group_laws_by_hierarchy(superior_laws)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_laws = sum(1 for laws in law_groups.values() for law in laws.values() if law is not None)
    current_idx = 0
    
    # 2단계: 각 그룹의 모든 계층 수집
    for base_name, laws in law_groups.items():
        group_content = {
            'base_name': base_name,
            'laws': {},
            'combined_articles': []
        }
        
        # 법률 → 시행령 → 시행규칙 순서로 수집
        for law_type in ['law', 'decree', 'rule']:
            law_name = laws[law_type]
            if law_name:
                current_idx += 1
                status_text.text(f"상위법령 조회 중... {law_name} ({current_idx}/{total_laws})")
                progress_bar.progress(current_idx / total_laws)
                
                law_content = get_superior_law_content(law_name)
                if law_content:
                    group_content['laws'][law_type] = law_content
                    # 새로운 데이터 구조 처리: content가 있으면 사용, articles가 있으면 변환
                    if 'content' in law_content:
                        # 연결된 본문이 있으면 그대로 저장
                        if 'combined_content' not in group_content:
                            group_content['combined_content'] = ""
                        group_content['combined_content'] += law_content['content'] + '\n'
                    elif 'articles' in law_content:
                        # 기존 articles 구조가 있으면 변환
                        group_content['combined_articles'].extend(law_content['articles'])
        
        if group_content['laws']:  # 하나 이상의 법령이 수집된 경우만 추가
            superior_laws_content.append(group_content)
    
    progress_bar.empty()
    status_text.empty()
    
    # 텍스트 길이 제한 (8만자) 및 관련성 필터링
    max_chars = 80000
    total_chars = 0
    
    # 각 법령 그룹의 텍스트 길이 계산
    for group in superior_laws_content:
        group_chars = 0
        
        # combined_content가 있는 경우
        if 'combined_content' in group and group['combined_content']:
            group_chars += len(group['combined_content'])
        
        # combined_articles가 있는 경우
        if 'combined_articles' in group and group['combined_articles']:
            for article in group['combined_articles']:
                group_chars += len(article.get('content', ''))
        
        # laws 구조가 있는 경우
        if 'laws' in group and group['laws']:
            for law_type, law_info in group['laws'].items():
                if law_info and 'articles' in law_info:
                    for article in law_info['articles']:
                        group_chars += len(article.get('content', ''))
        
        group['text_length'] = group_chars
        total_chars += group_chars
    
    return superior_laws_content

def chunk_text(text, chunk_size=1000, overlap=200):
    """텍스트를 청크로 분할하는 함수"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = text[start:end]
        
        # 문장 단위로 끝나도록 조정
        if end < text_length:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            last_break = max(last_period, last_newline)
            if last_break > start + chunk_size * 0.7:  # 너무 짧지 않으면 조정
                end = start + last_break + 1
                chunk = text[start:end]
        
        if chunk.strip():
            chunks.append({
                'text': chunk.strip(),
                'start': start,
                'end': end
            })
        
        start = end - overlap
    
    return chunks

def get_gemini_embedding(text, api_key):
    """Gemini를 사용하여 텍스트 임베딩 생성"""
    try:
        genai.configure(api_key=api_key)
        result = genai.embed_content(
            model="models/embedding-001",
            content=text,
            task_type="retrieval_document"
        )
        return result['embedding']
    except Exception as e:
        st.error(f"임베딩 생성 오류: {str(e)}")
        return None

def is_valid_text(text):
    """텍스트 품질 검사"""
    if not text or len(text.strip()) < 10:
        return False

    # 한글 깨짐 검사 (깨진 문자 비율이 30% 이상이면 제외)
    broken_chars = sum(1 for char in text if ord(char) > 55000)  # 한글 깨짐 문자 범위
    if len(text) > 0 and broken_chars / len(text) > 0.3:
        return False

    # 점선 과다 검사 (점선이 50% 이상이면 제외)
    dot_chars = text.count('·') + text.count('…') + text.count('.')
    if len(text) > 0 and dot_chars / len(text) > 0.5:
        return False

    # 반복 문자 과다 검사
    import re
    repeated_patterns = re.findall(r'(.)\1{10,}', text)  # 같은 문자가 10번 이상 반복
    if repeated_patterns:
        return False

    return True

def clean_text_content(text):
    """텍스트 정제"""
    import re

    # 1. 과도한 점선 제거
    text = re.sub(r'[·…]{3,}', ' ', text)
    text = re.sub(r'\.{3,}', ' ', text)

    # 2. 과도한 공백 정리
    text = re.sub(r'\s+', ' ', text)

    # 3. 페이지 번호 패턴 제거
    text = re.sub(r'\b\d+\s*페이지?\b', '', text)
    text = re.sub(r'\b\d+\s*쪽?\b', '', text)

    # 4. 목차 관련 패턴 제거
    text = re.sub(r'^[IVX]+\.?\s*', '', text, flags=re.MULTILINE)  # 로마숫자
    text = re.sub(r'^\d+\.?\s*$', '', text, flags=re.MULTILINE)   # 단독 숫자

    # 5. 반복되는 특수문자 정리
    text = re.sub(r'[~`!@#$%^&*()_+=\[\]{}|\\:";\'<>?/,-]{5,}', ' ', text)

    return text.strip()

def extract_legal_reasoning_from_analysis(analysis_text):
    """Gemini 분석 결과에서 법적 근거와 논리 추출"""
    import re

    extracted_context = {
        'legal_basis': [],      # 법적 근거 (법령, 조항)
        'reasoning': [],        # 추론 과정
        'key_concepts': [],     # 핵심 개념
        'problem_details': []   # 구체적인 문제점
    }

    # 1. 법령 및 조항 추출
    legal_references = re.findall(r'(?:지방자치법|헌법|행정기본법|건축법|도시계획법)\s*(?:제\s*\d+조?(?:의?\d+)?)?', analysis_text)
    extracted_context['legal_basis'].extend(legal_references)

    # 2. 법적 원칙/개념 추출
    legal_concepts = [
        '기관위임사무', '자치사무', '국가사무', '법률유보원칙', '권한배분',
        '상위법령', '법령우위', '조례제정권', '위임입법', '처분권한',
        '헌법위반', '기본권침해', '평등원칙', '비례원칙', '신뢰보호',
        '재산권침해', '영업의자유', '거주이전의자유', '표현의자유',
        '조세법률주의', '죄형법정주의', '적법절차', '정당한보상'
    ]

    for concept in legal_concepts:
        if concept in analysis_text:
            # 해당 개념 주변 문맥 추출 (앞뒤 50자)
            matches = re.finditer(re.escape(concept), analysis_text)
            for match in matches:
                start = max(0, match.start() - 50)
                end = min(len(analysis_text), match.end() + 50)
                context = analysis_text[start:end].strip()
                extracted_context['key_concepts'].append({
                    'concept': concept,
                    'context': context
                })

    # 3. 문제점 상세 내용 추출
    problem_patterns = [
        r'문제(?:점|가|는)[^.]*?(?:\.|$)',
        r'위법[^.]*?(?:\.|$)',
        r'위반[^.]*?(?:\.|$)',
        r'부적절[^.]*?(?:\.|$)',
        r'한계[^.]*?(?:\.|$)'
    ]

    for pattern in problem_patterns:
        matches = re.findall(pattern, analysis_text, re.DOTALL)
        extracted_context['problem_details'].extend(matches)

    # 4. 추론 과정 추출 (따라서, 그러므로, 왜냐하면 등)
    reasoning_patterns = [
        r'(?:따라서|그러므로|왜냐하면|이는|이에 따라)[^.]*?(?:\.|$)',
        r'(?:근거|이유|원인)는[^.]*?(?:\.|$)'
    ]

    for pattern in reasoning_patterns:
        matches = re.findall(pattern, analysis_text, re.DOTALL)
        extracted_context['reasoning'].extend(matches)

    return extracted_context


def detect_agency_delegation(superior_article: Dict, ordinance_article: Dict, source_type: str) -> Dict:
    """기관위임사무 특화 판별 함수"""
    
    superior_content = superior_article.get('content', '').lower()
    ordinance_content = ordinance_article.get('content', '').lower()
    
    # 1단계: 국가사무인지 판별
    national_affairs_indicators = [
        '건축허가', '개발행위허가', '환경영향평가', '도시계획',
        '산업단지', '관광단지', '택지개발', '도로개설',
        '하천점용', '산지전용', '농지전용', '산업입지',
        '국토계획', '지역계획', '광역계획'
    ]
    
    is_national_affair = any(indicator in superior_content for indicator in national_affairs_indicators)
    
    # 2단계: 지방자치단체 '장'에게 위임되었는지 확인
    delegation_to_head_indicators = [
        '시장', '군수', '구청장', '지방자치단체의 장',
        '시장이', '군수가', '구청장이', '장이',
        '위임한다', '위탁한다'
    ]
    
    is_delegated_to_head = any(indicator in superior_content for indicator in delegation_to_head_indicators)
    
    # 3단계: 조례가 해당 사무에 대해 별도 규정을 두고 있는지 확인
    ordinance_regulation_indicators = [
        '허가', '승인', '신고', '인가', '지정', '등록',
        '기준', '절차', '방법', '조건', '제한'
    ]
    
    has_ordinance_regulation = any(indicator in ordinance_content for indicator in ordinance_regulation_indicators)
    
    # 4단계: 위법성 판단
    is_agency_delegation = False
    severity = "낮음"
    evidence = []
    description = ""
    
    if is_national_affair and is_delegated_to_head and has_ordinance_regulation:
        is_agency_delegation = True
        severity = "매우 높음"
        description = "기관위임사무에 대해 조례로 별도 규정을 두어 지방자치법 제22조 위반"
        
        evidence.extend([
            f"국가사무 확인: {[ind for ind in national_affairs_indicators if ind in superior_content][:2]}",
            f"지방자치단체 장 위임 확인: {[ind for ind in delegation_to_head_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정 확인: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    elif is_national_affair and has_ordinance_regulation:
        # 국가사무인데 조례로 규정한 경우 (위임 대상 불확실)
        is_agency_delegation = True
        severity = "높음"
        description = "국가사무로 추정되는 사항에 대해 조례가 별도 규정, 기관위임사무 가능성"
        
        evidence.extend([
            f"국가사무 가능성: {[ind for ind in national_affairs_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    elif is_delegated_to_head and has_ordinance_regulation:
        # 지방자치단체 장 위임 + 조례 규정
        is_agency_delegation = True
        severity = "높음" 
        description = "지방자치단체 장에게 위임된 사무에 대해 조례로 별도 규정"
        
        evidence.extend([
            f"지방자치단체 장 위임: {[ind for ind in delegation_to_head_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    return {
        'is_agency_delegation': is_agency_delegation,
        'description': description,
        'evidence': evidence,
        'severity': severity,
        'national_affair': is_national_affair,
        'delegated_to_head': is_delegated_to_head,
        'has_regulation': has_ordinance_regulation
    }

def analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content):
    """조례와 상위법령 직접 비교 분석 함수 - 계층별 통합 검토"""
    analysis_results = []
    
    if not superior_laws_content:
        return "상위법령 정보가 없어 직접 비교 분석을 수행할 수 없습니다."

    # 조례에서 사무 관련 조문 추출
    ordinance_provisions = []
    lines = pdf_text.split('\n')
    current_article = ""
    current_content = ""
    
    for line in lines:
        line = line.strip()
        if line.startswith('제') and '조' in line:
            if current_article:
                ordinance_provisions.append({
                    'article': current_article,
                    'content': current_content.strip()
                })
            current_article = line
            current_content = ""
        else:
            current_content += line + " "
    
    # 마지막 조문 추가
    if current_article:
        ordinance_provisions.append({
            'article': current_article,
            'content': current_content.strip()
        })
    
    # 상위법령과 직접 비교 분석
    comparison_results = []
    
    for ordinance_provision in ordinance_provisions:
        if not ordinance_provision['content']:
            continue
            
        provision_analysis = {
            'ordinance_article': ordinance_provision['article'],
            'ordinance_content': ordinance_provision['content'],
            'superior_law_conflicts': [],
            'delegation_issues': [],
            'authority_issues': []
        }
        
        # 각 상위법령 그룹과 비교 (법률, 시행령, 시행규칙 통합)
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            
            # 연결된 본문이 있는 경우 간단한 키워드 매칭만 수행
            if 'combined_content' in law_group:
                superior_content_lower = law_group['combined_content'].lower()
                ordinance_lower = ordinance_provision['content'].lower()
                
                # 키워드 기반 관련성 확인
                common_keywords = []
                for word in ordinance_lower.split():
                    if len(word) > 2 and word in superior_content_lower:
                        common_keywords.append(word)
                
                if len(common_keywords) > 2:  # 최소 3개 이상의 공통 키워드가 있으면 관련성 있음
                    # 간단한 분석만 수행
                    continue
                else:
                    continue
            
            # 기존 방식 - articles가 있는 경우
            for superior_article in law_group.get('combined_articles', []):
                superior_content = superior_article['content'].lower()
                ordinance_lower = ordinance_provision['content'].lower()
                
                # 어느 계층(법률/시행령/시행규칙)에서 나온 조문인지 확인
                article_source = "법률"  # 기본값
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        for article in law_info['articles']:
                            if article['content'] == superior_article['content']:
                                if law_type == 'law':
                                    article_source = "법률"
                                elif law_type == 'decree':
                                    article_source = "시행령"
                                elif law_type == 'rule':
                                    article_source = "시행규칙"
                                break
                
                # 🆕 특화된 기관위임사무 판별 로직
                agency_delegation_result = detect_agency_delegation(
                    superior_article, ordinance_article, article_source
                )
                
                if agency_delegation_result['is_agency_delegation']:
                    provision_analysis['delegation_issues'].append({
                        'superior_law': f"{base_name} ({article_source})",
                        'superior_article': f"{superior_article['number']} {superior_article['title']}",
                        'superior_content': superior_article['content'],
                        'issue_type': '기관위임사무 위반',
                        'description': agency_delegation_result['description'],
                        'evidence': agency_delegation_result['evidence'],
                        'severity': agency_delegation_result['severity'],
                        'hierarchy': article_source
                    })
                
                # 직접적인 충돌 검사 - 계층별 위반 심각도 구분
                conflict_indicators = [
                    ('금지', '허용'), ('의무', '면제'), ('필수', '선택'),
                    ('강제', '임의'), ('반드시', '가능'), ('불가', '허용')
                ]
                
                for prohibit_word, allow_word in conflict_indicators:
                    if prohibit_word in superior_content and allow_word in ordinance_lower:
                        # 계층별 위반 심각도
                        severity = "심각" if article_source == "법률" else ("보통" if article_source == "시행령" else "경미")
                        
                        provision_analysis['superior_law_conflicts'].append({
                            'superior_law': f"{base_name} ({article_source})",
                            'superior_article': f"{superior_article['number']} {superior_article['title']}",
                            'conflict_type': f'{article_source} {prohibit_word} vs 조례 {allow_word}',
                            'superior_content': superior_article['content'],
                            'potential_violation': True,
                            'hierarchy': article_source,
                            'severity': severity
                        })
        
        if provision_analysis['delegation_issues'] or provision_analysis['superior_law_conflicts']:
            comparison_results.append(provision_analysis)
    
    return comparison_results

def create_analysis_prompt(pdf_text, search_results, superior_laws_content=None, relevant_guidelines=None, is_first_ordinance=False, comprehensive_analysis_results=None, theoretical_results=None):
    """분석 프롬프트 생성 함수"""
    prompt = (
        "=" * 80 + "\n"
        "⚠️ **최우선 지시사항: 검토 대상 조례 vs 참고자료 구분**\n"
        "=" * 80 + "\n"
        "당신은 조례 위법성 검토 전문가입니다.\n\n"
        "**핵심 임무**:\n"
        "1. [검토 대상 조례 원문] 섹션에 제시된 조례의 위법성을 검토\n"
        "2. [참고자료] 섹션의 자치법규 가이드라인/판례는 **위법 판단의 근거 자료일 뿐**\n"
        "3. **절대 참고자료를 검토 대상 조례로 착각하지 말 것**\n\n"
        "**분석 결과 형식**:\n"
        "- 반드시 아래 모든 섹션을 포함하여 답변:\n"
        "  1. [비교분석 요약표(조문별)]\n"
        "  2. [내 조례의 차별점 요약]\n"
        "  3. [검토 시 유의사항] - a) 소관사무의 원칙, b) 법률유보의 원칙, c) 법령우위의 원칙\n"
        "  4. [실무적 검토 포인트]\n\n"
        "=" * 80 + "\n\n"
        "🚨 **필독: 조례 위법 판단 법리적 가이드라인 (최우선 준수 사항)**\n\n"
        "**1. 조례 위법 판단의 원칙**\n"
        "- 조례는 자치입법권에 따라 폭넓은 재량을 가지며, 법률과 상충하거나 주민의 권리를 침해하지 않는 이상 위법으로 보지 않는다.\n"
        "- 조례는 법률유보, 포괄위임금지, 절차적 위법 등 요건이 엄격할 때만 위법성이 문제된다.\n"
        "- 단순한 표현 방식, 관행과의 차이, 타 지자체와의 비교만으로 위법 판단을 하지 않는다.\n\n"
        "**2. 자동해산·포괄조항 등 자치단체 조례에서 흔히 쓰는 규정은 원칙적으로 허용**\n"
        "- 특정안건형·비상설 위원회에 대한 자동 구성·자동 해산 규정은 조례가 정한 절차에 따른 것으로 보며 위법으로 판단하지 않는다.\n"
        "- '그 밖에 시장이 필요하다고 인정하는 경우' 등 보충적 포괄조항은 권리·의무를 제한하지 않는 한 자치입법권 범위에서 허용된다.\n"
        "- 이를 법률유보·포괄위임금지로 판단하지 않는다.\n\n"
        "**3. 자문위원회에 대해서는 '절차권', '구성권' 개념을 적용하지 않기**\n"
        "- 위원회는 자치단체장의 자문기구이며, 위원회 존폐는 주민의 절차권 또는 권익과 직접 관련되지 않는다.\n"
        "- 절차권 침해, 구성권 침해 등의 개념을 자문위원회에 적용하지 않는다.\n\n"
        "**4. 위법 판단은 반드시 '법률 상위규범과의 명확한 충돌'이 있을 때만 함**\n"
        "- 위법 판단을 할 때는 반드시 다음을 충족할 때만 위법성을 지적한다:\n"
        "  1) 조례가 상위법의 명령·금지를 명백히 위반하는 경우\n"
        "  2) 주민의 권리를 제한하거나 의무를 부과하면서 법적 근거가 명확히 없는 경우\n"
        "  3) 자치사무가 아닌 국가사무를 침해하는 경우\n"
        "- 그 외에는 '위법 가능성 있음'이라고 판단하지 않는다.\n\n"
        "---\n\n"
        "🚨 **중요 미션: 실제 위법 내용 찾기**\n"
        "너는 조례 위법성 전문 검토관이다. 일반적인 법리 설명이 아니라 **구체적인 위법 사항을 찾아내는 것**이 목표다.\n"
        "상위법령과 조례를 조문 대 조문으로 직접 비교하여 실제 충돌하는 부분을 찾아라.\n\n"
        "**검토 원칙:**\n"
        "- ❌ '이런 내용이 있으면 위법하다'는 일반론 금지\n"
        "- ✅ '조례 제3조는 도로교통법 제12조와 이렇게 충돌한다'는 구체적 지적 필수\n"
        "- ✅ 의심스러운 부분도 반드시 언급 (단, 위 가이드라인 1~4를 준수하여 신중히 판단)\n"
        "- ✅ 위법이 없으면 '위법 사항 없음'으로 명확히 결론\n\n"
        "=" * 80 + "\n"
        "📄 **[검토 대상 조례 원문 시작]**\n"
        "=" * 80 + "\n"
        "⚠️ **중요**: 아래 내용은 내가 업로드한 조례 PDF의 전체 내용이다.\n"
        "이 조례의 위법성을 검토하는 것이 당신의 임무이다.\n"
        "이 조례 이후에 제공되는 '참고자료', '가이드라인', '판례' 등은 모두 위법 판단의 근거일 뿐,\n"
        "검토 대상 조례 본문이 아니다. 절대 혼동하지 말 것.\n\n"
        f"{pdf_text}\n\n"
        "=" * 80 + "\n"
        "📄 **[검토 대상 조례 원문 종료]**\n"
        "=" * 80 + "\n\n"
    )
    
    # 상위법령 내용 추가 (계층별 그룹화)
    if superior_laws_content:
        prompt += "\n그리고 아래는 조례안에서 언급된 상위법령들의 실제 조문 내용이야. (법률, 시행령, 시행규칙을 계층별로 그룹화하여 통합 분석)\n"
        prompt += "---\n"
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            prompt += f"◆ {base_name}\n"
            
            # 연결된 본문이 있으면 사용
            if 'combined_content' in law_group:
                prompt += f"  본문 내용:\n{law_group['combined_content']}\n"
            else:
                # 기존 방식 - 각 계층별 법령 표시
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                        prompt += f"  [{type_name}] {law_info['law_name']}\n"
                
                # 통합된 조문 표시 (상위 15개만)
                prompt += f"  통합 조문 ({len(law_group['combined_articles'])}개):\n"
                for article in law_group['combined_articles'][:15]:  
                    prompt += f"    {article['number']} {article['title']}\n"
                    prompt += f"    {article['content']}\n\n"
        prompt += "---\n"
        
        # 상위법령 직접 비교 분석 결과 추가
        try:
            comparison_results = analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content)
            if comparison_results and isinstance(comparison_results, list) and len(comparison_results) > 0:
                prompt += "\n**중요: 조례와 상위법령 직접 비교 분석 결과**\n"
                prompt += "아래는 조례 조문과 상위법령을 하나씩 직접 비교한 결과이다. 이 분석을 바탕으로 기관위임사무 여부와 법령위반 가능성을 정확히 판단해줘.\n"
                prompt += "---\n"
                
                for result in comparison_results:
                    prompt += f"◆ {result['ordinance_article']}\n"
                    prompt += f"조례 내용: {result['ordinance_content'][:200]}...\n"
                    
                    if result['delegation_issues']:
                        prompt += "⚠️ 기관위임사무 가능성 발견:\n"
                        for issue in result['delegation_issues']:
                            prompt += f"  - {issue['superior_law']} {issue['superior_article']}\n"
                            prompt += f"    문제: {issue['description']}\n"
                    
                    if result['superior_law_conflicts']:
                        prompt += "🚨 상위법령 충돌 가능성 발견:\n"
                        for conflict in result['superior_law_conflicts']:
                            prompt += f"  - {conflict['superior_law']} {conflict['superior_article']}\n"
                            prompt += f"    충돌: {conflict['conflict_type']}\n"
                    
                    prompt += "\n"
                prompt += "---\n"
        except Exception as e:
            prompt += f"\n상위법령 직접 비교 분석 중 오류 발생: {str(e)}\n"
    
    # 자치법규 가이드라인 및 사례 추가
    if relevant_guidelines:
        prompt += "\n" + "=" * 80 + "\n"
        prompt += "📚 **[참고자료: 자치법규 가이드라인 및 위법 판단 기준]**\n"
        prompt += "=" * 80 + "\n"
        prompt += "⚠️ **중요 주의사항**: 아래 내용은 위법성 판단을 위한 **참고자료일 뿐**이다.\n"
        prompt += "**이것은 검토 대상 조례가 아니다.** 위에서 제시한 [검토 대상 조례 원문]과 혼동하지 말 것.\n"
        prompt += "아래는 자치법규 매뉴얼, 예전 위법 사례, 판례 등에서 검색된 관련 내용으로,\n"
        prompt += "위 조례의 위법성을 판단할 때 **근거 자료로만 활용**하라.\n\n"
        prompt += "**활용 방법**: 소관사무의 원칙, 법률유보의 원칙, 법령우위의 원칙 등 부분에 있어\n"
        prompt += "위 조례에서 조금이라도 문제가 될 것 같은 부분이 있다면,\n"
        prompt += "아래 자료에 수록된 예전에 문제가 되었던 사례와 검토 기준을 참조하여 판단하라.\n"
        prompt += "---\n"
        
        # 소스별로 그룹화하여 표시
        source_groups = {}
        for guideline in relevant_guidelines:
            source_store = guideline.get('source_store', '알 수 없는 자료')
            if source_store not in source_groups:
                source_groups[source_store] = []
            source_groups[source_store].append(guideline)
        
        for source_store, guidelines in source_groups.items():
            prompt += f"◆ 참고자료 출처: {source_store}\n"
            for i, guideline in enumerate(guidelines):
                similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                prompt += f"  [참고자료 {i+1}] (유사도: {similarity_score:.3f})\n"
                prompt += f"  {guideline['text']}\n\n"
        prompt += "---\n"
        prompt += "📚 **[참고자료 종료]**\n"
        prompt += "=" * 80 + "\n\n"
    
    # 종합 위법성 판례 분석 결과 추가
    if comprehensive_analysis_results and isinstance(comprehensive_analysis_results, list) and len(comprehensive_analysis_results) > 0:
        total_risks = sum(len(result['violation_risks']) for result in comprehensive_analysis_results)
        prompt += f"\n**🚨 중요: 종합 조례 위법성 판례 적용 결과 ({total_risks}개 위험)**\n"
        prompt += "참고 자료에서 검색된 실제 조례 위법 판례들(기관위임사무, 상위법령 위배, 법률유보 위배, 권한배분 위배 등)을\n"
        prompt += "현재 조례에 직접 적용한 분석 결과이다. 이 결과를 바탕으로 각 유형별 위법성을 정확히 판단하고 구체적인 개선방안을 제시해줘.\n"
        prompt += "---\n"
        
        for result in comprehensive_analysis_results:
            prompt += f"◆ {result['ordinance_article']}\n"
            prompt += f"조례 내용: {result['ordinance_content'][:150]}...\n"
            
            for i, risk in enumerate(result['violation_risks'][:2]):  # 상위 2개만 포함
                prompt += f"  위험 {i+1}: {risk['violation_type']} (위험도: {risk['risk_score']:.2f}/1.0)\n"
                prompt += f"  관련 판례: {risk['case_summary'][:150]}...\n"
                if risk['legal_principle'] != "해당없음":
                    prompt += f"  법적 원칙: {risk['legal_principle']}\n"
                prompt += f"  개선 권고: {risk['recommendation']}\n"
                prompt += f"  판례 출처: {risk['case_source']}\n\n"
            
            if len(result['violation_risks']) > 2:
                prompt += f"  ...외 {len(result['violation_risks']) - 2}개 추가 위험\n\n"
        prompt += "---\n"

    # 🆕 검색된 관련 판례/이론 추가
    if theoretical_results and isinstance(theoretical_results, list) and len(theoretical_results) > 0:
        prompt += f"\n**📚 중요: 발견된 문제점 관련 판례/이론 ({len(theoretical_results)}개)**\n"
        prompt += "이는 1차 분석에서 발견된 문제점들과 직접 관련된 판례와 법리이다.\n"
        prompt += "아래 판례들을 참고하여 현재 조례의 위법성을 정확히 판단하고 구체적인 개선방안을 제시해줘.\n"
        prompt += "---\n"

        for i, theory in enumerate(theoretical_results[:5]):  # 상위 5개만 포함
            context_rel = theory.get('context_relevance', 0)
            matched_concepts = theory.get('matched_concepts', [])
            similarity = theory.get('similarity', 0)

            prompt += f"◆ 관련 판례/이론 {i+1} (관련도: {context_rel:.2f}, 유사도: {similarity:.2f})\n"
            if matched_concepts:
                prompt += f"관련 개념: {', '.join(matched_concepts)}\n"

            # 내용 미리보기 (300자로 제한)
            content = theory.get('content', theory.get('text', '내용 없음'))
            content_preview = content[:300] + "..." if len(content) > 300 else content
            prompt += f"내용: {content_preview}\n\n"

        prompt += "**⚠️ 중요**: 위 판례들은 조례의 문제점과 직접 관련이 있으므로, 이를 근거로 현재 조례의 위법성을 구체적으로 지적하고 개선방안을 제시하라.\n"
        prompt += "---\n"

    if is_first_ordinance:
        prompt += (
            "※ 참고: 이 조례는 17개 시도 중 최초로 제정되는 조례로, 타시도 조례가 존재하지 않습니다.\n"
            "타시도 조례가 없는 상황에서, 아래 기준에 따라 조례의 적정성, 상위법령과의 관계, 실무적 검토 포인트 등을 중심으로 분석해줘.\n"
        )
    else:
        prompt += "그리고 아래는 타시도 조례명과 각 조문 내용이야.\n"
        for result in search_results:
            prompt += f"조례명: {result['name']}\n"
            for idx, article in enumerate(result['content']):
                prompt += f"제{idx+1}조: {article}\n"
    
    prompt += (
        "---\n"
        "아래 기준에 따라 분석해줘. 반드시 한글로 답변해줘.\n"
        "1. [비교분석 요약표(조문별)]\n"
        "- 표의 컬럼: 조문(내 조례), 주요 내용, 타 시도 유사 조항, 동일 여부, 차이 및 내 조례 특징, 추천 조문\n"
        "- 반드시 내 조례(PDF로 업로드한 조례)의 조문만을 기준으로, 각 조문별로 타 시도 조례와 비교해 표로 정리(내 조례에 없는 조문은 비교하지 말 것)\n"
        "- '추천 조문' 칸에는 타 시도 조례와 비교해 무난하게 생각되는 조문 예시를 한글로 작성\n\n"
        "2. [내 조례의 차별점 요약] (별도 소제목)\n"
        "- 타 시도 조례와 비교해 독특하거나 구조적으로 다른 점, 내 조례만의 관리/운영 방식 등 요약\n\n"
        "3. [검토 시 유의사항] (별도 소제목)\n"
        "각 항목마다 일반인도 이해할 수 있도록 쉬운 말로 부연설명도 함께 작성해줘.\n"
        "다음 원칙들을 기준으로 검토해줘:\n"
        "a) 소관사무의 원칙 - *기관위임사무는 조례 제정 금지**\n"
        " 자치사무의 예시는 지방자치법 제13조제2항에 열거 되어 있음**\n"
        " 개별 법령에서 국가 또는 중앙행정기관의 장을 권한 주체로 정하고 있는 경우 국가사무로 보아야 함. 국가사무에 관한 사항을 규정한 조례는 위법. 국가사무 여부를 판단함에 있어서 지방자치법 제15조를 고려할 수 있음. 다만, 법령에서 일정 사항을 조례로 정할 수 있다고 규정한다면 그 사무가 국가사무나 자치사무 관계없이 조례 제정 가능**\n"
        " 지방자치단체 또는 지방자치단체의 장을 권한주체로 정하고 있는 경우 자치사무로 보아야 함**\n"
        " 법령에 국가와 지방자치단체를 사무 수행의 주체로 병렬적으로 규정하는 경우 국가사무와 자치사무 성질을 모두 가지므로 조례로 규율 가능능**\n"
        "**기관위임사무 정의**: 국가사무를 지방자치단체장(특별시장, 도지사, 광역시장, 시장, 군수, 구청장)에게 위임한 사무, 조례에서 위임한게 아니고 법률, 시행령, 시행규칙에서 위임한 것을 말함\n"
        "**핵심 원칙**: 기관위임사무에 대해서는 조례 제정이 원칙적으로 금지됨 (지방자치법 제22조)\n"
        "**판별 기준**: \n"
        "  1) 사무가 국가사무인지 확인 (예: 건축허가, 도시계획, 환경영향평가 등)\n"
        "  2) 해당 사무가 지방자치단체장(특별시장, 도지사, 광역시장, 시장, 군수, 구청장)에게 위임되었는지 확인\n"
        "  3) 자치사무인지 기관위임사무인지 판단함에 있어 법령의 규정형식과 취지를 우선 고려해야 할것이나 그 외에도 사무의 성질이 전국적으로 통일적인 처리를 요구하는 사무인지 경비부담과 최종적인 책임귀속 주체등도 고려해 판단\n"
        "**위법 사례**: 건축허가, 개발행위허가, 환경영향평가 등 국가위임사무에 대해 조례로 추가 규정을 둔 경우\n"
        "- 지방자치단체의 자치사무와 법령에 의해 위임된 단체위임사무에 대해서만 제정 가능한지\n"
        "- 사무의 성격이 전국적으로 통일적 처리를 요구하는지 여부 검토\n\n"
        "b) 법률 유보의 원칙\n"
        "- 주민의 권리를 제한하거나 의무를 부과에 관한 사항이나 벌칙을 정할 때에는 법률의 위임이 있어야 함\n"
        "- 상위 법령에서 위임받지 않은 권한을 행사하는지 확인 (단, 권리·의무를 제한하지 않는 조직·절차 규정은 제외)\n"
        "- 상위 법령의 위임 범위를 명백히 초과하는지 검토\n\n"
        "**⚠️ 중요: 포괄조항에 대한 올바른 판단 기준**\n"
        "- '그 밖에 시장이 필요하다고 인정하는 경우' 등 보충적 포괄조항은:\n"
        "  ① 주민의 권리·의무를 직접 제한하지 않고\n"
        "  ② 자치단체의 자문·심의기구 운영에 관한 사항이며\n"
        "  ③ 다른 지자체 조례에서도 흔히 사용되는 경우\n"
        "  → **원칙적으로 적법한 규정으로 판단**\n"
        "- 타 지자체 조례에 유사 조항이 다수 존재한다면 이는 조례 관행으로 인정되므로 위법으로 보지 않음\n"
        "- 법률유보 위반으로 판단하기 위해서는 '주민의 권리 제한 또는 의무 부과'라는 요건이 반드시 충족되어야 함\n\n"
        "c) 법령우위의 원칙 위반 여부\n"
        "- **조례가 법령에 위반되는지 여부는 법령과 조례의 각각의 규정 취지, 규정의 목적과 내용 및 효과 등을 비교하여 양자 사이에 모순, 저촉이 있는지 여부에 따라 개별적, 구체적으로 결정해야 함**\n"
        "- **일반론이 아닌 구체적 충돌 지점을 찾을 것 - 단순히 '다르다'는 것만으로는 위법이 아님**\n"
        "- '다른 조례에 특별 규정이 없으면 본 조례가 우선'이라는 규정은 다른 조례와 비교했을 때 우선한다는 것이지 상위법령보다 우선한다는 것이 아니기 때문에 적법함\n"
        "- 위에 제시된 상위법령 본문을 한 조문씩 꼼꼼히 읽고 조례와 대조할 것\n\n"
        "**⚠️ 중요: 자치단체 위원회·자문기구 운영 규정에 대한 판단 기준**\n"
        "- 자치단체가 설치하는 각종 위원회, 자문기구, 협의체 등의 구성·운영에 관한 사항은:\n"
        "  ① 자치단체의 내부 조직·절차에 관한 사항으로서\n"
        "  ② 주민의 권리·의무와 직접 관련이 없고\n"
        "  ③ 자치입법권의 핵심 영역에 해당하므로\n"
        "  → **상위법에 명시적 금지 규정이 없는 한 원칙적으로 적법**\n"
        "- 특정안건형 위원회의 자동 구성·자동 해산 조항은 조례 제정권의 범위 내에서 허용됨\n"
        "- '절차권 침해', '구성권 침해' 등의 개념은 주민의 권익과 직접 관련된 경우에만 적용되며, 자문기구에는 적용하지 않음\n\n"
        "**검토 방법**:\n"
        "1) 조례 제1조부터 마지막 조문까지 하나씩 검토\n"
        "2) 각 조례 조문의 내용과 관련된 상위법령 조문을 찾아서 직접 비교\n"
        "3) 다음과 같은 구체적 충돌이 있는지 확인:\n"
        "   - 조례에서 규율하는 내용에 관한 법령이 없는 경우 평등의 원칙, 비례의 원칙, 명확성의 원칙 같은 법의 일반원칙에 위반되지 않는 지 검토\n"
        "   - 조례의 목적과 취지가 법령의 목적과 취지와 같은 경우에도 법령의 취지가 전국에 걸쳐 일률적인 규율을 하려는 것이 아니라 각 지자체가 지방 실정에 맞게 별도로 규율하는 것을 용인한다고 해설될때는 법령에 위반 되는 것이 아님 \n"
        "   - 수익적 내용이면 법령에 근거가 없어도 조례로 정할 수 있으므로 법령에서 조례로 다르게 정할 수 없다고 규정하지 않는 이상 법령과 다르게 조례에 규정할 수 있는 여지가 많음\n"
        "   - 침익적 내용이면 법률에서 위임받은 범위에서만 조례로 정할 수 있으므로 법령과 다르게 조례에 규정할 수 있는 여지가 거의 없음\n"
        "   - 조례가 상위법령보다 강한 의무나 제재를 부과하는 경우\n"
        "   - 조례가 상위법령의 위임 범위를 명백히 벗어나는 경우\n"
        "   - 조례가 상위법령에서 국가나 중앙행정기관 소관으로 정한 사무에 관여하는 경우\n\n"
        "**위법 발견 시 반드시 다음 형식으로 구체적으로 명시:**\n"
        "  🚨 **위법 사항 발견** (상위 가이드라인 1~4를 충족하는 경우에만 지적)\n"
        "  * **조례 조문**: 제○조 ○항 - \"조례의 정확한 문구\"\n"
        "  * **상위법령**: ○○법 제○조 ○항 - \"상위법령의 정확한 문구\"\n"
        "  * **충돌 내용**: 구체적으로 어떤 부분이 어떻게 위배되는지 상세 설명 (추상적 설명 금지)\n"
        "  * **위법 유형**: (법령우위 위반/법률유보 위반/기관위임사무 위반)\n"
        "  * **위법 판단 근거**: 위 가이드라인 4조 중 어느 요건을 충족하는지 명시 (1)상위법 명령·금지 위반, 2)권리제한·의무부과 근거 부재, 3)국가사무 침해)\n"
        "  * **개선 방안**: 상위법령에 맞는 구체적 수정안\n\n"
        "**⚠️ 중요: 위법 판단의 엄격성**\n"
        "- 위법 사항이 없는 경우 '위법 사항을 발견하지 못했음'으로 명확히 결론\n"
        "- 의심스러운 부분이 있더라도 **위 가이드라인 1~4를 충족하지 않으면 위법으로 판단하지 않음**\n"
        "- 단순 표현 차이, 타 지자체와의 조문 구성 차이만으로는 위법성을 지적하지 않음\n"
        "- 포괄조항, 자동해산 조항, 자문기구 운영 조항 등 자치입법권 범위 내 사항은 위법으로 보지 않음\n\n"
        "4. 실무적 검토 포인트\n"
        "- 조례의 집행 과정에서 발생할 수 있는 문제점\n"
        "- 개선이 필요한 부분과 그 방향성\n\n"
    )

    # 상위법령별 개별 위반 여부 검토 (Gemini 전용 프롬프트 추가)
    if superior_laws_content:
        prompt += "\n5. [상위법령별 개별 위반 여부 검토]\n"
        prompt += "위에서 제시한 상위법령들 각각에 대해 개별적으로 다음 기준에 따라 상세 분석해줘:\n\n"

        section_num = 1
        for law_group in superior_laws_content:
            base_name = law_group['base_name']

            prompt += f"5-{section_num}) [{base_name} 위반 여부 검토]\n"
            prompt += f"상위법령명: {base_name}\n"

            # 해당 법령의 본문 일부 재참조
            if 'combined_content' in law_group:
                law_content_preview = law_group['combined_content'][:2000]
                prompt += f"상위 법령 본문 일부:\n{law_content_preview}\n\n"
            elif 'combined_articles' in law_group and law_group['combined_articles']:
                prompt += "상위 법령 주요 조문:\n"
                for article in law_group['combined_articles'][:5]:  # 처음 5개 조문만
                    prompt += f"  {article['number']} {article['title']}\n"
                    prompt += f"  {article['content'][:300]}...\n\n"

            prompt += f"**🔍 {base_name} 세부 검토 지시사항: (⚠️ 상위 법리적 가이드라인 1~4 준수)**\n"
            prompt += "위 상위법령 본문을 조례와 한 조문씩 직접 대조하여 다음을 수행하라:\n\n"
            prompt += "  ① **조문별 직접 대조 분석** (단순 차이는 위법이 아님)\n"
            prompt += f"  - 조례의 각 조문이 {base_name}의 어떤 조문과 관련되는지 식별\n"
            prompt += f"  - {base_name}에서 금지/허용/의무화하는 사항과 조례 내용 직접 비교\n"
            prompt += "  - **명백한 상충**이 있을 때만 지적 (단순 표현 차이, 조문 구성 차이는 제외)\n\n"
            prompt += "  ② **권한 범위 초과 여부** (자치입법권 범위 고려)\n"
            prompt += f"  - {base_name}에서 국가/중앙행정기관 전담으로 '명시적'으로 정한 사무가 있는지 확인\n"
            prompt += "  - 조례가 해당 사무에 개입하고 있는지 점검 (단, 위임 규정이 있으면 적법)\n"
            prompt += "  - 위임 범위를 '명백히' 벗어난 규정이 있는지 확인 (해석상 여지가 있으면 위법 아님)\n\n"
            prompt += "  ③ **구체적 위법 사항 발견 시** (⚠️ 가이드라인 4조 요건 충족 시에만)\n"
            prompt += "  🚨 **위법 발견 보고 형식:**\n"
            prompt += "  * **문제 조문**: 조례 제○조 - \"정확한 조문 내용\"\n"
            prompt += f"  * **관련 상위법령**: {base_name} 제○조 - \"정확한 조문 내용\"\n"
            prompt += "  * **위법 사유**: 구체적인 충돌/위반 내용 (추상적 설명 금지)\n"
            prompt += "  * **가이드라인 4조 충족 여부**: (1)상위법 명령·금지 위반 / (2)권리제한·의무부과 근거 부재 / (3)국가사무 침해 중 해당 항목 명시\n"
            prompt += "  * **위법 심각도**: 경미/보통/심각\n"
            prompt += "  * **수정 방안**: 구체적인 개선 방향\n\n"
            prompt += "  ④ **위법이 아닌 경우 명확히 기재**\n"
            prompt += "  - 단순 표현 차이, 조문 구성 차이는 '위법 아님'으로 명시\n"
            prompt += "  - 포괄조항, 자동해산 조항, 자문기구 운영 조항 등은 '자치입법권 범위 내 적법'으로 판단\n"
            prompt += "  - 의심 사항이 있더라도 가이드라인 1~4를 충족하지 않으면 '위법 아님'으로 결론\n\n"

            section_num += 1

    return prompt

def parse_table_from_text(text_content):
    """텍스트에서 표 형태의 내용을 파싱하여 Word 표 데이터로 변환"""
    tables_data = []
    lines = text_content.split('\n')
    current_table = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 표의 시작을 감지 (|가 포함된 라인)
        if '|' in line and len([cell for cell in line.split('|') if cell.strip()]) >= 3:
            # 표 헤더인지 구분 (첫 번째 |로 시작하는 라인)
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]

            if current_table is None:
                # 새 표 시작
                current_table = {'headers': cells, 'rows': []}
                tables_data.append(current_table)
            else:
                # 구분선이 아닌 데이터 행인지 확인
                if not all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                    current_table['rows'].append(cells)
        else:
            # 표가 끝남
            if current_table is not None:
                current_table = None

    return tables_data

def add_table_to_doc(doc, table_data):
    """Word 문서에 표 추가"""
    if not table_data['headers']:
        return

    # 열 수 계산
    max_cols = len(table_data['headers'])
    for row in table_data['rows']:
        max_cols = max(max_cols, len(row))

    # 행 수 계산 (헤더 + 데이터 행)
    row_count = 1 + len(table_data['rows'])

    if row_count == 1:  # 헤더만 있는 경우 스킵
        return

    # 표 생성
    table = doc.add_table(rows=row_count, cols=max_cols)
    table.style = 'Table Grid'
    table.autofit = True

    # 헤더 추가
    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data['headers']):
        if i < len(header_cells):
            header_cells[i].text = header
            # 헤더 스타일링
            paragraph = header_cells[i].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.bold = True

    # 데이터 행 추가
    for row_idx, row_data in enumerate(table_data['rows']):
        if row_idx + 1 < len(table.rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = cell_data

def create_comparison_document(pdf_text, search_results, analysis_results, superior_laws_content=None, relevant_guidelines=None):
    """비교 분석 문서 생성 함수"""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)

    # 제목 추가
    title = doc.add_heading('조례 비교 분석 결과', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'분석 일시: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')

    # 상위법령 정보 추가 (계층별 그룹화)
    if superior_laws_content:
        doc.add_heading('검토된 상위법령', level=2)
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            
            # 그룹 제목 추가
            doc.add_paragraph(f"◆ {base_name}")
            
            # 연결된 본문이 있는 경우
            if 'combined_content' in law_group:
                content_length = len(law_group['combined_content'])
                doc.add_paragraph(f"  • 본문 {content_length:,}자")
            else:
                # 기존 방식 - 각 계층별 법령 정보 표시
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                        doc.add_paragraph(f"  • {law_info['law_name']} ({type_name}) - {len(law_info['articles'])}개 조문")
                
                combined_articles = law_group.get('combined_articles', [])
                doc.add_paragraph(f"  총 {len(combined_articles)}개 조문 통합 검토")
            
            doc.add_paragraph("")
        doc.add_paragraph("")
    
    # 활용된 자치법규 자료 정보 추가
    if relevant_guidelines:
        doc.add_heading('활용된 자치법규 참고자료', level=2)
        
        # 소스별로 그룹화
        source_groups = {}
        for guideline in relevant_guidelines:
            source_store = guideline.get('source_store', '알 수 없는 자료')
            if source_store not in source_groups:
                source_groups[source_store] = []
            source_groups[source_store].append(guideline)
        
        for source_store, guidelines in source_groups.items():
            doc.add_paragraph(f"◆ {source_store} ({len(guidelines)}개 내용)")
            for i, guideline in enumerate(guidelines):
                similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                doc.add_paragraph(f"   • 내용 {i+1} (유사도: {similarity_score:.3f})")
        doc.add_paragraph("")

    # 최종 분석 결과만 추가 (중복 방지)
    # 우선순위: 자료 참고 보강분석 > OpenAI 추가 분석 > 1차 분석
    final_report = None
    for result in reversed(analysis_results):  # 역순으로 최신 결과 우선
        if 'error' not in result:
            if "자료 참고 보강분석" in result.get('model', ''):
                final_report = result
                break
            elif "자료 참고" in result.get('model', '') or "OpenAI" in result.get('model', ''):
                final_report = result
                break

    # 자료 참고나 OpenAI가 없으면 1차 분석 사용
    if not final_report:
        for result in analysis_results:
            if 'error' not in result and "1차 분석" in result.get('model', ''):
                final_report = result
                break

    # 최종 보고서가 있으면 추가
    if final_report:
        doc.add_heading(f'📋 {final_report["model"]}', level=2)
        content = final_report.get('content') or final_report.get('analysis', '')

        # 🆕 표 파싱 및 처리
        tables_data = parse_table_from_text(content)

        # 텍스트를 섹션별로 처리
        lines = content.split('\n')
        current_section = []

        for line in lines:
            line = line.strip()

            # 표 라인인지 확인 (|가 포함된 라인)
            if '|' in line and len([cell for cell in line.split('|') if cell.strip()]) >= 3:
                # 표 시작 전까지의 텍스트 처리
                if current_section:
                    for text_line in current_section:
                        text_line_clean = text_line.strip()
                        if text_line_clean:
                            # 제목 라인 처리 (1., 2., 3. 등으로 시작하거나 [로 시작하는 경우)
                            if (text_line_clean.startswith(('1.', '2.', '3.', '4.', '5.')) or
                                text_line_clean.startswith('[') and text_line_clean.endswith(']')):
                                # 마크다운 기호 제거하고 제목으로 추가
                                title_text = re.sub(r'[#*`>\-\[\]]+', '', text_line_clean)
                                doc.add_heading(title_text, level=3)
                            else:
                                # 일반 텍스트 - 마크다운 볼드(**text**) 처리
                                if clean_text := text_line_clean.strip():
                                    p = doc.add_paragraph()
                                    # **텍스트** 형식의 볼드 처리
                                    parts = re.split(r'(\*\*[^*]+\*\*)', clean_text)
                                    for part in parts:
                                        if part.startswith('**') and part.endswith('**'):
                                            # 볼드 텍스트
                                            run = p.add_run(part[2:-2])
                                            run.bold = True
                                        elif part:
                                            # 일반 텍스트 (남은 마크다운 기호 제거)
                                            clean_part = re.sub(r'[#`>]+', '', part)
                                            if clean_part:
                                                p.add_run(clean_part)
                    current_section = []

                # 표 처리는 skip (이미 tables_data에서 처리됨)
                continue
            else:
                # 구분선이 아닌 경우만 텍스트로 추가
                if not (line.replace('-', '').replace(':', '').replace('|', '').strip() == ''):
                    current_section.append(line)

        # 마지막 섹션 처리
        if current_section:
            for text_line in current_section:
                text_line_clean = text_line.strip()
                if text_line_clean:
                    if (text_line_clean.startswith(('1.', '2.', '3.', '4.', '5.')) or
                        text_line_clean.startswith('[') and text_line_clean.endswith(']')):
                        title_text = re.sub(r'[#*`>\-\[\]]+', '', text_line_clean)
                        doc.add_heading(title_text, level=3)
                    else:
                        # 일반 텍스트 - 마크다운 볼드(**text**) 처리
                        if clean_text := text_line_clean.strip():
                            p = doc.add_paragraph()
                            # **텍스트** 형식의 볼드 처리
                            parts = re.split(r'(\*\*[^*]+\*\*)', clean_text)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    # 볼드 텍스트
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                elif part:
                                    # 일반 텍스트 (남은 마크다운 기호 제거)
                                    clean_part = re.sub(r'[#`>]+', '', part)
                                    if clean_part:
                                        p.add_run(clean_part)

        # 🆕 파싱된 표들을 Word 문서에 추가
        for table_data in tables_data:
            add_table_to_doc(doc, table_data)
            doc.add_paragraph("")  # 표 간격
    else:
        # 최종 보고서가 없으면 오류 표시
        doc.add_heading('⚠️ 분석 결과 없음', level=2)
        doc.add_paragraph('분석 결과를 생성할 수 없습니다.')

        # 오류 메시지 추가
        for result in analysis_results:
            if 'error' in result:
                doc.add_paragraph(f"❌ {result['model']} 오류: {result['error']}")

    return doc

def send_error_report(subject, body, attachment_data=None, attachment_name=None):
    """이메일 전송 함수"""
    try:
        email_config = _email_config_from_secrets()
        sender_email = email_config.get("sender_email")
        sender_password = email_config.get("sender_password")
        receiver_email = "lsh4676@korea.kr"

        if not sender_email or not sender_password:
            st.error("이메일 설정이 올바르지 않습니다. .streamlit/secrets.toml을 확인해주세요.")
            return False

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        if attachment_data and attachment_name:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment_data)
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {attachment_name}',
            )
            msg.attach(part)

        # SMTP 서버 연결 (Gmail 예시)
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        
        return True
    except Exception as e:
        st.error(f"이메일 전송 중 오류 발생: {str(e)}")
        return False

def main():
    # 헤더
    st.markdown("""
    <div class="main-header">
        <h1>🏛️ 광역지자체 조례 검색, 비교, 분석</h1>
        <p>17개 광역지자체의 조례를 검색하고, AI를 활용하여 비교 분석할 수 있는 도구입니다.</p>
    </div>
    """, unsafe_allow_html=True)

    # 사이드바
    with st.sidebar:
        st.header("📋 작업 순서")
        st.markdown("""
        <div class="step-card">
            <strong>1단계:</strong> 조례 검색 및 Word 저장<br>
            검색어를 입력하여 17개 시도의 조례를 검색하고 3단 비교 형태로 MS Word 문서를 생성합니다.
        </div>
        <div class="step-card">
            <strong>2단계:</strong> 조례안 PDF 업로드<br>
            제정 또는 개정할 조례안 PDF 파일을 업로드합니다.
        </div>
        <div class="step-card">
            <strong>3단계:</strong> AI 비교 분석<br>
            업로드한 조례안과 타 시도 조례를 AI로 비교 분석하여 MS Word 문서를 생성합니다.
        </div>
        """, unsafe_allow_html=True)

        st.header("🤖 AI 분석 엔진")

        # Ollama Cloud 상태 확인
        ollama_available = bool(st.session_state.ollama_api_key and st.session_state.ollama_api_key != "YOUR_OLLAMA_API_KEY_HERE")

        if ollama_available:
            st.success("✅ **무료 AI 분석 서비스 활성화됨**")
            st.info("🚀 API 키 입력 없이 바로 분석을 시작할 수 있습니다!")
            use_ollama = st.checkbox(
                "Ollama Cloud 사용 (무료, 권장)",
                value=st.session_state.use_ollama_cloud,
                help="120B 파라미터의 고성능 AI 모델을 무료로 사용합니다. API 키 발급이 필요 없습니다."
            )
            st.session_state.use_ollama_cloud = use_ollama
        else:
            st.warning("⚠️ Ollama Cloud 서비스가 설정되지 않았습니다.")
            use_ollama = False
            st.session_state.use_ollama_cloud = False

        st.markdown("---")

        # 고급 설정 (선택적)
        with st.expander("⚙️ 고급 설정 (선택사항)", expanded=False):
            st.markdown("**추가 AI 서비스** (선택적으로 사용)")
            st.caption("배포 시 Cloudtype 환경 변수 `GEMINI_API_KEY` 또는 `GOOGLE_API_KEY`, `OPENAI_API_KEY`로도 설정할 수 있습니다.")
            st.text_input(
                "Gemini API 키",
                type="password",
                key="gemini_api_key_input",
                help="Google AI Studio 키. 환경 변수 GEMINI_API_KEY 또는 GOOGLE_API_KEY와 동일하면 자동 반영됩니다.",
            )
            st.text_input(
                "OpenAI API 키",
                type="password",
                key="openai_api_key_input",
                help="OpenAI API 키. 환경 변수 OPENAI_API_KEY와 동일하면 자동 반영됩니다.",
            )
            gemini_api_key = st.session_state.gemini_api_key_input
            openai_api_key = st.session_state.openai_api_key_input

            # Gemini File Search Store Manager 초기화
            if gemini_api_key and st.session_state.gemini_store_manager is None:
                try:
                    st.session_state.gemini_store_manager = get_gemini_store_manager(gemini_api_key)
                    st.success("✅ Gemini File Search 초기화 완료")
                except Exception as e:
                    st.warning(f"⚠️ Gemini File Search 초기화 실패: {e}")

            st.markdown("---")
            st.subheader("🔍 검색 엔진 설정")

            use_gemini = st.checkbox(
                "Gemini File Search 사용",
                value=st.session_state.use_gemini_search if gemini_api_key else False,
                help="기존 방식 대신 Gemini File Search API를 사용합니다. 더 정확한 검색 결과를 제공합니다.",
                disabled=not gemini_api_key
            )
            st.session_state.use_gemini_search = use_gemini

            if use_gemini:
                if st.session_state.gemini_store_manager:
                    st.success("✅ Gemini File Search 활성화됨")
                else:
                    st.warning("⚠️ Gemini API 키를 먼저 입력해주세요")

        st.header("ℹ️ 서비스 안내")
        st.markdown("""
        <div class="step-card">
            <strong>🎉 무료 AI 분석 서비스</strong><br>
            본 서비스는 Ollama Cloud의 고성능 AI 모델(120B 파라미터)을 무료로 제공합니다.<br>
            <strong>API 키 발급 없이 바로 사용 가능합니다!</strong>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("📋 추가 AI 서비스 안내 (선택사항)", expanded=False):
            st.markdown("""
            더 다양한 분석이 필요한 경우, 아래 서비스를 추가로 사용할 수 있습니다.

            ### 🤖 Gemini API (선택사항)
            - **용도**: Gemini File Search를 통한 정밀 검색
            - **발급**: [aistudio.google.com](https://aistudio.google.com)
            - **무료 할당량**: 월 1,000번 요청

            ### 🧠 OpenAI API (선택사항)
            - **용도**: 추가 교차 검증 분석
            - **발급**: [platform.openai.com](https://platform.openai.com)
            - **요금**: 사용량 기반 과금

            ⚠️ **참고**: 추가 API 키 없이도 기본 분석은 완전히 작동합니다!
            """)


    # 메인 컨텐츠
    tab1, tab2, tab3 = st.tabs(["1️⃣ 조례 검색", "2️⃣ PDF 업로드", "3️⃣ AI 분석"])

    with tab1:
        st.header("조례 검색")
        
        # 검색 폼 (Enter 키 지원)
        with st.form(key="search_form"):
            col1, col2 = st.columns([3, 1])
            with col1:
                search_query = st.text_input(
                    "검색어를 입력하세요 (키워드)", 
                    placeholder="예: 청년지원 (Enter 키로도 검색 가능)", 
                    value=st.session_state.search_query,
                    help="검색어를 입력한 후 Enter 키를 누르거나 검색 버튼을 클릭하세요."
                )
            with col2:
                search_button = st.form_submit_button("🔍 검색", type="primary")

        # 검색 실행 (Enter 키 또는 버튼 클릭 시)
        if search_button and search_query.strip():
            st.session_state.search_query = search_query.strip()
            st.session_state.word_doc_ready = False  # 문서 준비 상태 초기화
            st.session_state.selected_ordinances = []  # 선택된 조례 초기화
            
            with st.spinner("검색 중... 잠시만 기다려주세요."):
                try:
                    results, total_count = search_ordinances(search_query.strip())
                    st.session_state.search_results = results
                    # 초기에는 모든 조례를 선택된 상태로 설정
                    st.session_state.selected_ordinances = list(range(len(results)))
                    st.success(f"검색 완료! 총 {len(results)}건의 조례가 검색되었습니다.")
                except Exception as e:
                    st.error(f"검색 중 오류 발생: {str(e)}")
                    st.session_state.search_results = []

        # 검색 결과가 있을 때 조례 선택 및 Word 문서 생성 기능
        if st.session_state.search_results:
            results = st.session_state.search_results
            
            # 검색 결과 요약 표시
            if not st.session_state.word_doc_ready:
                st.success(f"검색 완료! 총 {len(results)}건의 조례가 검색되었습니다.")
            
            # 조례 선택 섹션
            st.subheader("📋 Word 문서에 포함할 조례 선택")
            
            # 전체 선택/해제 버튼
            col1, col2, col3 = st.columns([1, 1, 2])
            
            with col1:
                if st.button("✅ 전체 선택", key="select_all_btn"):
                    st.session_state.selected_ordinances = list(range(len(results)))
                    for idx in range(len(results)):
                        st.session_state[f"ordinance_checkbox_{idx}"] = True
                    st.rerun()
            
            with col2:
                if st.button("❌ 전체 해제", key="deselect_all_btn"):
                    st.session_state.selected_ordinances = []
                    for idx in range(len(results)):
                        st.session_state[f"ordinance_checkbox_{idx}"] = False
                    st.rerun()
            
            with col3:
                selected_count = len(st.session_state.selected_ordinances)
                st.markdown(f"**선택된 조례: {selected_count}개 / 총 {len(results)}개**")
            
            # 조례 선택 체크박스
            st.markdown("---")
            
            # 조례별 체크박스 표시
            for idx, result in enumerate(results):
                # 🆕 단순화: 체크박스 상태를 직접 관리
                is_selected = idx in st.session_state.selected_ordinances
                checkbox_key = f"ordinance_checkbox_{idx}"

                # 체크박스와 조례명을 한 줄에 표시
                current_checked = st.checkbox(
                    f"**{result['metro']}** - {result['name']}",
                    value=is_selected,
                    key=checkbox_key
                )

                # 🆕 상태 변경 감지 및 즉시 반영
                if current_checked != is_selected:
                    if current_checked:
                        # 체크됨 - 목록에 추가
                        if idx not in st.session_state.selected_ordinances:
                            st.session_state.selected_ordinances.append(idx)
                    else:
                        # 체크 해제됨 - 목록에서 제거
                        if idx in st.session_state.selected_ordinances:
                            st.session_state.selected_ordinances.remove(idx)
            
            st.markdown("---")
            
            # Word 문서 생성 버튼
            col1, col2 = st.columns([1, 1])
            
            with col1:
                # 선택된 조례가 있을 때만 생성 버튼 활성화
                disabled = len(st.session_state.selected_ordinances) == 0
                
                if st.button("📄 선택된 조례로 Word 문서 생성", type="secondary", key="create_word_btn", disabled=disabled):
                    if st.session_state.selected_ordinances:
                        try:
                            with st.spinner("Word 문서 생성 중..."):
                                # 선택된 조례만 필터링
                                selected_results = [results[i] for i in st.session_state.selected_ordinances]
                                
                                # Word 문서 생성
                                doc = create_word_document(st.session_state.search_query, selected_results)
                                
                                # Word 문서를 바이트로 변환
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                doc_bytes = doc_io.getvalue()
                                
                                # 세션 상태에 저장
                                st.session_state.word_doc_data = doc_bytes
                                st.session_state.word_doc_ready = True
                                
                            st.success(f"✅ 선택된 {len(selected_results)}개 조례로 Word 문서가 생성되었습니다!")
                            st.rerun()  # 페이지 새로고침으로 다운로드 버튼 표시
                            
                        except Exception as e:
                            st.error(f"❌ Word 문서 생성 중 오류 발생: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())
                    else:
                        st.warning("조례를 하나 이상 선택해주세요.")
                
                if disabled:
                    st.caption("⚠️ 조례를 하나 이상 선택해주세요.")
            
            with col2:
                # Word 문서가 준비되면 다운로드 버튼 표시
                if st.session_state.word_doc_ready and st.session_state.word_doc_data:
                    filename = f"조례_검색결과_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.download_button(
                        label="💾 Word 문서 다운로드",
                        data=st.session_state.word_doc_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_word_btn"
                    )
            # 상세 검색 결과 표시 (조례 내용 확인용)
            st.subheader("📖 조례 내용 상세보기")
            
            for idx, result in enumerate(results):
                # 🆕 단순화: 선택 상태만 텍스트로 표시
                is_selected = idx in st.session_state.selected_ordinances
                status = " ✅ 선택됨" if is_selected else " ⭕ 선택안됨"

                with st.expander(f"{result['metro']} - {result['name']}{status}", expanded=False):
                    st.markdown(f"<div class='metro-name'>{result['metro']}</div>", unsafe_allow_html=True)
                    st.markdown(f"<div class='law-title'>{result['name']}</div>", unsafe_allow_html=True)
                    
                    if result['content']:
                        for article_idx, article in enumerate(result['content']):
                            st.markdown(f"**제{article_idx+1}조**")
                            st.markdown(article)
                            st.markdown("---")
                    else:
                        st.markdown("*(조문 없음)*")
        
        elif search_button and not search_query.strip():
            st.error("검색어를 입력해주세요.")
        elif not st.session_state.search_results:
            st.info("검색어를 입력하고 Enter 키를 누르거나 검색 버튼을 클릭하세요.")

    with tab2:
        st.header("조례안 PDF 업로드")
        
        uploaded_file = st.file_uploader("제정 또는 개정할 조례안 PDF 파일을 업로드하세요", type=['pdf'])
        
        if uploaded_file is not None:
            st.session_state.uploaded_pdf = uploaded_file
            st.success(f"파일이 업로드되었습니다: {uploaded_file.name}")
            
            # PDF 내용 미리보기 - expander로 변경하여 재실행 방지
            with st.expander("PDF 내용 미리보기", expanded=False):
                with st.spinner("PDF 내용을 읽는 중..."):
                    pdf_text = extract_pdf_text(uploaded_file)
                    if pdf_text:
                        st.text_area("PDF 내용", pdf_text[:2000] + "..." if len(pdf_text) > 2000 else pdf_text, height=300)
                    else:
                        st.error("PDF 내용을 읽을 수 없습니다.")

    with tab3:
        st.header("AI 비교 분석")

        # 조건 확인 - PDF가 업로드되고 AI 서비스가 사용 가능하면 분석 가능
        pdf_uploaded = st.session_state.uploaded_pdf is not None
        has_ollama = st.session_state.use_ollama_cloud and bool(st.session_state.ollama_api_key and st.session_state.ollama_api_key != "YOUR_OLLAMA_API_KEY_HERE")
        has_api_key = bool(gemini_api_key or openai_api_key) or has_ollama
        has_search_results = bool(st.session_state.search_results)

        if not pdf_uploaded:
            st.warning("📄 먼저 PDF 파일을 업로드해주세요.")
        elif not has_api_key:
            st.warning("🔑 AI 분석 서비스가 설정되지 않았습니다. 관리자에게 문의하세요.")
        else:
            # 검색 결과 여부에 따라 안내 메시지 표시
            if not has_search_results:
                st.info("💡 **최초 제정 조례 분석**")
                st.markdown("""
                검색된 타 시도 조례가 없습니다. 이는 다음과 같은 경우일 수 있습니다:
                - 🆕 **최초 제정 조례**: 17개 시도 중 최초로 제정되는 조례
                - 🔍 **검색어 불일치**: 다른 키워드로 재검색 후 분석 권장
                
                검색 결과가 없어도 조례안의 **법적 검토**와 **상위법령 위반 여부** 분석이 가능합니다.
                """)
            else:
                st.success(f"📊 {len(st.session_state.search_results)}개의 타 시도 조례와 비교 분석합니다.")
        
        # 분석 가능한 조건일 때 분석 인터페이스 표시
        if pdf_uploaded and has_api_key:
            # 검색어 입력 (선택사항)
            search_query_analysis = st.text_input(
                "검색어 (분석용)", 
                value=st.session_state.search_query if st.session_state.search_query else "", 
                key="analysis_query",
                help="검색어를 입력하면 더 정확한 분석이 가능합니다. (선택사항)"
            )
            
            # 분석 타입 표시 (선택된 조례 수 반영)
            if not has_search_results:
                analysis_type = "최초 제정 조례 분석"
            elif hasattr(st.session_state, 'selected_ordinances') and st.session_state.selected_ordinances:
                selected_count = len(st.session_state.selected_ordinances)
                analysis_type = f"선택된 {selected_count}개 타 시도 조례와 비교 분석"
            else:
                analysis_type = f"전체 {len(st.session_state.search_results)}개 타 시도 조례와 비교 분석"
            st.markdown(f"**분석 유형**: {analysis_type}")
            
            # 자동 참고 자료 검색 옵션 (문제 발견 시 자동 활용)
            use_auto_search = st.checkbox(
                "🔍 문제 발견 시 자동 참고 자료 검색",
                value=True,
                help="법적 문제점을 발견한 경우 Gemini File Search를 통해 자동으로 관련 판례 및 법령 자료를 검색하여 근거를 보강합니다."
            )
            
            # 🆕 저장된 분석 결과가 있으면 먼저 표시
            if hasattr(st.session_state, 'analysis_results') and st.session_state.analysis_results:
                st.info("💾 **이전 분석 결과가 저장되어 있습니다**")

                # 메타데이터 표시
                if hasattr(st.session_state, 'analysis_metadata'):
                    metadata = st.session_state.analysis_metadata
                    st.caption(f"📅 분석 시간: {metadata.get('analysis_timestamp', '알 수 없음')}")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📋 이전 분석 결과 보기", use_container_width=True):
                        st.session_state.show_previous_analysis = True
                        st.rerun()
                with col2:
                    if st.button("🔄 새로 분석하기", use_container_width=True):
                        # 기존 결과 초기화
                        if hasattr(st.session_state, 'analysis_results'):
                            del st.session_state.analysis_results
                        if hasattr(st.session_state, 'analysis_metadata'):
                            del st.session_state.analysis_metadata
                        if hasattr(st.session_state, 'show_previous_analysis'):
                            del st.session_state.show_previous_analysis
                        st.rerun()

            # 이전 분석 결과 표시
            if hasattr(st.session_state, 'show_previous_analysis') and st.session_state.show_previous_analysis and hasattr(st.session_state, 'analysis_results'):
                analysis_results = st.session_state.analysis_results
                metadata = st.session_state.analysis_metadata

                st.markdown("---")
                st.subheader("📋 저장된 AI 분석 결과")

                # 분석 완료 메시지 (저장된 메타데이터 기반)
                has_problems = metadata.get('has_problems', False)
                relevant_guidelines = metadata.get('relevant_guidelines')
                loaded_stores = metadata.get('loaded_stores')
                is_first_ordinance = metadata.get('is_first_ordinance', False)

                if has_problems and relevant_guidelines and loaded_stores:
                    st.success(f"🎯 **복합 자료 보강 분석 완료**: 문제점 탐지 → {len(loaded_stores)}개 자료 참고 → 보강 분석")
                elif has_problems and relevant_guidelines:
                    st.success("🎯 **지능형 분석 완료**: 문제점 탐지 → 자료 검색 → 보강 분석")
                elif has_problems:
                    st.info("⚠️ **문제점 탐지 분석 완료**: 자료 검색 없이 기본 분석만 수행")
                else:
                    st.success("✅ **기본 분석 완료**: 특별한 문제점이 발견되지 않음")

                # 분석 결과 요약
                analysis_count = len([r for r in analysis_results if 'error' not in r])
                if analysis_count > 0:
                    # 🆕 저장된 메타데이터에서 선택된 조례 수 반영
                    if is_first_ordinance:
                        analysis_type_text = "최초 제정 조례"
                    else:
                        saved_search_results = metadata.get('search_results_for_analysis', [])
                        selected_count = len(saved_search_results)
                        analysis_type_text = f"선택된 {selected_count}개 타 시도 조례 비교"
                    st.markdown(f"**📋 분석 유형**: {analysis_type_text}")
                    st.markdown(f"**🤖 수행된 분석**: {analysis_count}개")
                    if relevant_guidelines:
                        guideline_count = len(relevant_guidelines) if isinstance(relevant_guidelines, list) else 0
                        st.markdown(f"**📚 참고 가이드라인**: {guideline_count}개")

                # 분석 결과 표시
                for result in analysis_results:
                    if 'error' not in result:
                        final_report = result
                        # 모델에 따른 구분 표시
                        if "보강" in final_report['model']:
                            st.success("🎯 **복합 자료 참고 보강 분석 결과**")
                            st.caption(f"📚 **활용 모델**: {final_report['model']}")
                        elif "자료 참고" in final_report['model']:
                            st.success("🎯 **참고 자료 기반 보강 분석 결과**")
                        elif "OpenAI" in final_report['model']:
                            st.info("📊 **OpenAI 추가 분석 결과**")
                        else:
                            st.info("🤖 **Gemini 기본 분석 결과**")
                        # 보고서 내용
                        st.markdown(final_report['content'])

                # 오류 메시지 표시
                for result in analysis_results:
                    if 'error' in result:
                        st.error(f"❌ {result['model']} 오류: {result['error']}")

                # Word 문서 다운로드 (메타데이터에서 복원)
                with st.spinner("저장된 분석 결과 Word 문서 생성 중..."):
                    superior_laws_content = metadata.get('superior_laws_content')
                    search_results_for_analysis = metadata.get('search_results_for_analysis')
                    pdf_text = metadata.get('pdf_text')
                    doc = create_comparison_document(pdf_text, search_results_for_analysis, analysis_results, superior_laws_content, relevant_guidelines)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_bytes = doc_io.getvalue()
                    # 파일명 설정
                    if has_problems and relevant_guidelines and loaded_stores:
                        stores_count = len(loaded_stores)
                        filename_prefix = f"복합자료보강분석({stores_count}개자료)" if is_first_ordinance else f"조례비교_복합자료분석({stores_count}개자료)"
                    elif has_problems and relevant_guidelines:
                        filename_prefix = "자료참고보강분석" if is_first_ordinance else "조례비교_자료분석"
                    elif has_problems:
                        filename_prefix = "문제점탐지분석" if is_first_ordinance else "조례비교_문제점분석"
                    else:
                        filename_prefix = "최초조례_기본분석" if is_first_ordinance else "조례_기본비교분석"
                    st.download_button(
                        label="📄 분석 결과 Word 문서 다운로드",
                        data=doc_bytes,
                        file_name=f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_previous_analysis"
                    )

                st.markdown("---")
                st.markdown("💡 **새로 분석하려면 위의 '🔄 새로 분석하기' 버튼을 클릭하세요.**")

            else:
                # 저장된 결과가 없거나 새 분석을 선택한 경우만 분석 시작 버튼 표시
                # 🆕 선택된 조례가 없는 경우 경고 표시
                if has_search_results and hasattr(st.session_state, 'selected_ordinances') and not st.session_state.selected_ordinances:
                    st.warning("⚠️ 비교할 조례가 선택되지 않았습니다. 조례 검색 탭에서 조례를 선택하거나, 선택 없이 최초 제정 조례 분석을 진행하세요.")

                if st.button("🤖 AI 분석 시작", type="primary"):
                    with st.spinner("AI가 조례를 분석하고 있습니다... 잠시만 기다려주세요."):
                        # PDF 텍스트 추출
                        pdf_text = extract_pdf_text(st.session_state.uploaded_pdf)
                    
                    if not pdf_text:
                        st.error("PDF 텍스트를 읽을 수 없습니다.")
                    else:
                        # 1단계: 상위법령 추출
                        with st.spinner("조례안에서 상위법령을 추출하고 있습니다..."):
                            superior_laws = extract_superior_laws(pdf_text)
                        
                        if superior_laws:
                            # 2단계: 상위법령 내용 조회
                            with st.spinner("국가법령정보센터에서 상위법령 내용을 조회하고 있습니다..."):
                                superior_laws_content = get_all_superior_laws_content(superior_laws)
                            
                            if superior_laws_content:
                                # 상위법령 조회 성공 (디버그 메시지 제거)
                                pass

                                # 🆕 상위법령 본문 내용 디버깅 표시 (expander로 변경하여 재실행 방지)
                                with st.expander("🔍 Gemini가 참조할 상위법령 본문 내용 미리보기", expanded=False):
                                    for i, law_group in enumerate(superior_laws_content):
                                        st.markdown(f"### [{i+1}] {law_group['base_name']}")

                                        # 연결된 본문이 있는 경우
                                        if 'combined_content' in law_group and law_group['combined_content']:
                                            content = law_group['combined_content']
                                            st.markdown(f"**본문 길이**: {len(content):,}자")
                                            st.text_area(
                                                f"{law_group['base_name']} 본문",
                                                content,
                                                height=200,
                                                key=f"content_{i}"
                                            )
                                        else:
                                            # 개별 법령별 표시
                                            for law_type, law_info in law_group['laws'].items():
                                                if law_info and 'articles' in law_info:
                                                    type_name = {"law": "법률", "decree": "시행령", "rule": "시행규칙"}[law_type]
                                                    st.markdown(f"#### {type_name}")

                                                    # 조문별 내용 표시 (처음 5개만)
                                                    for j, article in enumerate(law_info['articles'][:5]):
                                                        st.markdown(f"**제{article.get('number', '?')}조** {article.get('title', '')}")
                                                        content = article.get('content', '')[:500]
                                                        st.markdown(f"```\n{content}{'...' if len(article.get('content', '')) > 500 else ''}\n```")

                                                    if len(law_info['articles']) > 5:
                                                        st.markdown(f"... (총 {len(law_info['articles'])}개 조문 중 5개만 표시)")

                                        st.markdown("---")
                                
                                # 2-1단계: 상위법령 직접 비교 분석
                                try:
                                    comparison_results = analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content)
                                    
                                    if comparison_results and isinstance(comparison_results, list) and len(comparison_results) > 0:
                                        st.warning(f"⚠️ {len(comparison_results)}개 조문에서 잠재적 문제점이 발견되었습니다!")
                                        
                                        with st.expander("🔍 상위법령 직접 비교 분석 결과", expanded=True):
                                            for i, result in enumerate(comparison_results):
                                                st.markdown(f"**🔍 {result['ordinance_article']}**")
                                                st.markdown(f"조례 내용: {result['ordinance_content'][:300]}...")
                                                
                                                if result['delegation_issues']:
                                                    st.error("⚠️ **기관위임사무 가능성 발견**")
                                                    for issue in result['delegation_issues']:
                                                        st.markdown(f"- **관련 상위법령**: {issue['superior_law']} {issue['superior_article']}")
                                                        st.markdown(f"- **문제점**: {issue['description']}")
                                                        st.markdown(f"- **상위법령 내용**: {issue['superior_content'][:200]}...")
                                                
                                                if result['superior_law_conflicts']:
                                                    st.error("🚨 **상위법령 충돌 가능성 발견**")
                                                    for conflict in result['superior_law_conflicts']:
                                                        st.markdown(f"- **관련 상위법령**: {conflict['superior_law']} {conflict['superior_article']}")
                                                        st.markdown(f"- **충돌 유형**: {conflict['conflict_type']}")
                                                        st.markdown(f"- **상위법령 내용**: {conflict['superior_content'][:200]}...")
                                                
                                                st.markdown("---")
                                    else:
                                        st.success("✅ 상위법령 직접 비교에서 명백한 충돌이나 기관위임사무 문제를 발견하지 못했습니다.")
                                        
                                except Exception as e:
                                    st.error(f"상위법령 직접 비교 분석 중 오류: {str(e)}")
                                
                                # 상위법령 내용 미리보기 (계층별 그룹화) - expander로 변경하여 재실행 방지
                                with st.expander("🔍 조회된 상위법령 내용 미리보기 (계층별)", expanded=False):
                                    for idx, law_group in enumerate(superior_laws_content):
                                        base_name = law_group['base_name']

                                        # 연결된 본문이 있는 경우
                                        if 'combined_content' in law_group:
                                            content_preview = law_group['combined_content'][:500] + "..." if len(law_group['combined_content']) > 500 else law_group['combined_content']
                                            with st.expander(f"📋 {base_name} ({len(law_group['combined_content']):,}자)", expanded=False):
                                                st.text_area("본문 내용", content_preview, height=300, disabled=True, key=f"content_{base_name}_{idx}")
                                        else:
                                            # 기존 방식
                                            with st.expander(f"📋 {base_name} 계층 ({len(law_group.get('combined_articles', []))}개 조문)", expanded=False):

                                                # 계층별 법령 정보 표시
                                                st.markdown("**📚 포함된 법령:**")
                                                for law_type, law_info in law_group['laws'].items():
                                                    if law_info and 'articles' in law_info:
                                                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                                                        st.markdown(f"- [{type_name}] {law_info['law_name']} ({len(law_info['articles'])}개 조문)")

                                                st.markdown("\n**📖 통합 조문 (처음 5개):**")
                                                combined_articles = law_group.get('combined_articles', [])
                                                for article in combined_articles[:5]:
                                                    st.markdown(f"**{article['number']} {article['title']}**")
                                                    st.markdown(article['content'][:200] + "..." if len(article['content']) > 200 else article['content'])
                                                    st.markdown("---")
                                                if len(combined_articles) > 5:
                                                    st.markdown(f"*(총 {len(combined_articles)}개 조문 중 5개만 표시)*")
                            else:
                                st.warning("⚠️ 상위법령 내용 조회에 실패했습니다. 일반적인 분석을 진행합니다.")
                        else:
                            st.info("ℹ️ 조례안에서 명시적인 상위법령을 찾을 수 없습니다.")
                            superior_laws_content = None
                        
                        # 3단계: Gemini 1차 분석 (문제점 탐지)
                        analysis_results = []
                        is_first_ordinance = not has_search_results

                        # 🆕 선택된 조례만 분석에 사용
                        if has_search_results and hasattr(st.session_state, 'selected_ordinances'):
                            selected_results = [st.session_state.search_results[i] for i in st.session_state.selected_ordinances if i < len(st.session_state.search_results)]
                            search_results_for_analysis = selected_results
                            st.info(f"📋 선택된 {len(search_results_for_analysis)}개 조례로 분석을 진행합니다.")
                        else:
                            search_results_for_analysis = st.session_state.search_results if has_search_results else []

                        # 🆕 3-1단계: 위법 판례 선제 검색 (모든 조문에 대해)
                        theoretical_results = []
                        if st.session_state.gemini_store_manager and gemini_api_key:
                            with st.spinner("📚 업로드된 조례의 모든 조문에 대한 위법 판례를 검색하고 있습니다..."):
                                try:
                                    # PDF에서 조례명 추출 (처음 10줄에서)
                                    ordinance_name = ""
                                    # re 모듈은 파일 상단에서 이미 import됨
                                    lines = pdf_text.split('\n')
                                    for line in lines[:10]:
                                        line = line.strip()
                                        # 조례명 패턴: "○○시 ○○ 조례" 또는 "○○에 관한 조례"
                                        name_match = re.search(r'([\w가-힣]+(?:시|도|군|구)\s+[\w가-힣\s]+(?:조례|조례안))', line)
                                        if not name_match:
                                            name_match = re.search(r'([\w가-힣\s]+에\s+관한\s+조례(?:안)?)', line)
                                        if name_match:
                                            ordinance_name = name_match.group(1).strip()
                                            st.info(f"📋 조례명: {ordinance_name}")
                                            break

                                    # 조례에서 모든 조문 추출
                                    ordinance_articles = []
                                    current_article = ""
                                    current_content = ""

                                    # 조례명을 첫 번째 항목으로 추가 (검색에 활용)
                                    if ordinance_name:
                                        ordinance_articles.append(f"조례명: {ordinance_name}")

                                    for line in lines:
                                        line = line.strip()
                                        if line.startswith('제') and '조' in line:
                                            if current_article and current_content:
                                                ordinance_articles.append(f"{current_article} {current_content.strip()}")
                                            current_article = line
                                            current_content = ""
                                        else:
                                            current_content += line + " "

                                    # 마지막 조문 추가
                                    if current_article and current_content:
                                        ordinance_articles.append(f"{current_article} {current_content.strip()}")

                                    # ❌ 조기 검색 제거: 조문만으로는 맥락이 부족하여 RAG 효과가 낮음
                                    # 대신 1차 분석 후 분석 결과를 바탕으로 정밀 검색 수행 (2480행 참조)
                                    # if ordinance_articles:
                                    #     theoretical_results_raw = search_violation_cases_gemini(
                                    #         ordinance_articles=ordinance_articles,
                                    #         api_key=gemini_api_key,
                                    #         store_manager=st.session_state.gemini_store_manager,
                                    #         max_results=12
                                    #     )
                                    #     theoretical_results = theoretical_results_raw
                                    #
                                    #     if theoretical_results:
                                    #         st.success(f"✅ {len(theoretical_results)}개의 관련 위법 판례/재의제소 사례를 찾았습니다!")

                                    # 초기화: 1차 분석 후 재검색으로 채워질 예정
                                    theoretical_results = []

                                    # 세션에 저장하여 프롬프트에서 사용 (나중에 재검색으로 업데이트됨)
                                    st.session_state.theoretical_results = theoretical_results

                                    # 미리보기 제거: 조기 검색을 제거했으므로 이 시점에는 비어있음
                                    # 1차 분석 후 정밀 검색 결과는 2580행의 "정밀 검색 결과 미리보기"에서 표시됨

                                except Exception as e:
                                    st.warning(f"⚠️ 조문 추출 중 오류 (계속 진행): {str(e)}")
                                    theoretical_results = []
                                    st.session_state.theoretical_results = theoretical_results

                        # AI 1차 분석 (문제점 탐지용) - Ollama Cloud 우선 사용
                        first_analysis = None
                        has_problems = False
                        analysis_model_name = ""
                        rag_context = ""

                        # Ollama Cloud를 우선적으로 사용
                        if has_ollama:
                            try:
                                # comprehensive_analysis_results 초기화
                                comprehensive_analysis_results = None

                                # RAG 벡터스토어 로드 및 검색
                                with st.spinner("📚 자치법규 매뉴얼 및 판례 자료를 로드하고 있습니다..."):
                                    vectorstores = load_rag_vectorstores()

                                if vectorstores:
                                    # 조례명 가져오기 (이전에 정의되었는지 확인)
                                    current_ordinance_name = st.session_state.get('current_ordinance_name', '')
                                    if not current_ordinance_name:
                                        # PDF 텍스트에서 조례명 추출 시도
                                        name_match = re.search(r'([가-힣\s]+(?:조례|규칙))', pdf_text[:500])
                                        if name_match:
                                            current_ordinance_name = name_match.group(1).strip()

                                    # 조례 내용에서 잠재적 위법성 키워드 추출
                                    potential_issues = []
                                    ordinance_sample = pdf_text[:3000]  # 처음 3000자 분석

                                    # 위법성 관련 패턴 감지
                                    issue_patterns = {
                                        '수수료': ['수수료', '사용료', '요금', '부담금'],
                                        '벌칙': ['벌칙', '과태료', '과징금', '벌금', '제재'],
                                        '권리제한': ['제한', '금지', '의무', '허가', '신고', '등록'],
                                        '재정': ['지원', '보조금', '출연', '예산', '재정'],
                                        '조직': ['위원회', '협의회', '심의회', '기구', '조직'],
                                        '인사': ['임명', '위촉', '해임', '겸직', '자격'],
                                        '위임': ['위임', '대행', '위탁', '대리'],
                                        '주민권리': ['주민', '청구', '투표', '참여', '공개']
                                    }

                                    for issue_type, keywords in issue_patterns.items():
                                        for keyword in keywords:
                                            if keyword in ordinance_sample:
                                                potential_issues.append(issue_type)
                                                break

                                    # 중복 제거
                                    potential_issues = list(set(potential_issues))

                                    # 조례 내용 기반 동적 검색 쿼리 생성
                                    if potential_issues:
                                        issue_keywords = ' '.join(potential_issues[:3])  # 최대 3개 이슈
                                        search_query = f"{current_ordinance_name} {issue_keywords} 조례 위법 판단"
                                        st.info(f"🔍 감지된 잠재적 검토 필요 사항: {', '.join(potential_issues)}")
                                    else:
                                        search_query = f"{current_ordinance_name} 조례 위법 판단 기준 자치사무"

                                    rag_results = search_rag_context(search_query, vectorstores, top_k=5)

                                    if rag_results:
                                        st.success(f"✅ {len(rag_results)}개의 관련 자치법규 자료를 찾았습니다!")

                                        # RAG 컨텍스트 구성
                                        rag_context = "\n\n[참고 자료: 자치법규 매뉴얼 및 재의·제소 판례]\n"
                                        for i, result in enumerate(rag_results[:5], 1):
                                            source_name = "자치법규 매뉴얼" if result['source'] == 'manual' else "재의·제소 판례"
                                            rag_context += f"\n--- {source_name} 참고자료 {i} ---\n"
                                            rag_context += result['text'][:1500] + "\n"

                                        rag_context += "\n[중요] 위 참고 자료를 바탕으로 실제 위법 여부를 신중하게 판단하세요. 단순히 상위법과 다르다고 해서 위법한 것이 아닙니다. 자치사무와 위임사무를 구분하고, 지방자치단체의 조례제정권 범위를 고려하세요.\n"

                                        with st.expander("📖 RAG 검색 결과 미리보기", expanded=False):
                                            for i, result in enumerate(rag_results[:5], 1):
                                                source_name = "자치법규 매뉴얼" if result['source'] == 'manual' else "재의·제소 판례"
                                                st.markdown(f"**{i}. {source_name}** (점수: {result.get('score', 0)})")
                                                st.text(result['text'][:500] + "...")
                                                st.markdown("---")
                                    else:
                                        st.info("RAG 검색에서 관련 자료를 찾지 못했습니다.")

                                # 1차 분석용 프롬프트 생성 (RAG 컨텍스트 포함)
                                theoretical_results = st.session_state.get('theoretical_results', None)
                                first_prompt = create_analysis_prompt(pdf_text, search_results_for_analysis, superior_laws_content, None, is_first_ordinance, comprehensive_analysis_results, theoretical_results)

                                # RAG 컨텍스트를 프롬프트 앞부분에 추가
                                if rag_context:
                                    first_prompt = rag_context + "\n\n" + first_prompt

                                # Ollama Cloud 전송 프롬프트 디버깅 표시
                                with st.expander("🔍 AI에게 전송되는 프롬프트 내용 확인", expanded=False):
                                    st.markdown("### 프롬프트 구조 분석")
                                    st.markdown(f"**전체 길이**: {len(first_prompt):,}자")
                                    st.markdown(f"**사용 모델**: Ollama Cloud (gpt-oss:120b)")
                                    if rag_context:
                                        st.markdown(f"**RAG 컨텍스트 포함**: ✅ ({len(rag_context):,}자)")

                                    # 전체 프롬프트 표시 (처음 2000자만)
                                    st.text_area(
                                        "전체 프롬프트 (처음 2000자)",
                                        first_prompt[:2000] + "..." if len(first_prompt) > 2000 else first_prompt,
                                        height=400,
                                        key="full_prompt_ollama"
                                    )

                                with st.spinner("🤖 Ollama Cloud AI가 조례를 분석하고 있습니다..."):
                                    response_text = call_ollama_cloud_api(first_prompt)

                                if response_text:
                                    first_analysis = response_text
                                    analysis_model_name = "Ollama Cloud (gpt-oss:120b)"

                                    # 문제점 키워드 탐지
                                    problem_keywords = [
                                        "위반", "문제", "충돌", "부적절", "개선", "수정", "보완",
                                        "법령 위반", "상위법령", "위법", "불일치", "모순", "우려"
                                    ]

                                    has_problems = any(keyword in first_analysis for keyword in problem_keywords)

                                    if has_problems:
                                        st.warning(f"⚠️ AI가 잠재적 문제점을 발견했습니다!")

                                    # 분석 결과 저장
                                    analysis_results.append({
                                        'model': 'Ollama Cloud (1차 분석)',
                                        'analysis': first_analysis,
                                        'has_problems': has_problems
                                    })
                                else:
                                    st.error("Ollama Cloud 1차 분석 응답이 비어있습니다.")

                            except Exception as e:
                                st.error(f"Ollama Cloud 1차 분석 오류: {str(e)}")
                                analysis_results.append({
                                    'model': 'Ollama Cloud (1차 분석)',
                                    'analysis': '',
                                    'error': str(e)
                                })

                        # Gemini API가 있으면 추가 분석 (선택적)
                        elif gemini_api_key:
                            try:
                                # comprehensive_analysis_results 초기화
                                comprehensive_analysis_results = None

                                genai.configure(api_key=gemini_api_key)
                                model = genai.GenerativeModel('gemini-2.0-flash-lite')

                                # 1차 분석용 프롬프트 (문제점 탐지 중심)
                                # 검색된 판례 정보 가져오기
                                theoretical_results = st.session_state.get('theoretical_results', None)
                                first_prompt = create_analysis_prompt(pdf_text, search_results_for_analysis, superior_laws_content, None, is_first_ordinance, comprehensive_analysis_results, theoretical_results)

                                # 🆕 Gemini 전송 프롬프트 디버깅 표시 - expander로 변경하여 재실행 방지
                                with st.expander("🔍 Gemini에게 전송되는 프롬프트 내용 확인", expanded=False):
                                    st.markdown("### 프롬프트 구조 분석")
                                    st.markdown(f"**전체 길이**: {len(first_prompt):,}자")

                                    # 상위법령 내용 부분만 추출
                                    if "상위법령들의 실제 조문 내용" in first_prompt:
                                        law_start = first_prompt.find("상위법령들의 실제 조문 내용")
                                        law_end = first_prompt.find("3. [검토 시 유의사항]")
                                        if law_end == -1:
                                            law_end = law_start + 5000  # 기본값

                                        law_content = first_prompt[law_start:law_end]
                                        st.markdown(f"**상위법령 내용 길이**: {len(law_content):,}자")

                                        st.text_area(
                                            "상위법령 관련 프롬프트 내용",
                                            law_content[:3000] + "..." if len(law_content) > 3000 else law_content,
                                            height=300,
                                            key="prompt_law_content"
                                        )

                                    # 전체 프롬프트 표시 (처음 2000자만)
                                    st.text_area(
                                        "전체 프롬프트 (처음 2000자)",
                                        first_prompt[:2000] + "..." if len(first_prompt) > 2000 else first_prompt,
                                        height=400,
                                        key="full_prompt"
                                    )
                                
                                response = model.generate_content(first_prompt)
                                
                                if response and hasattr(response, 'text') and response.text:
                                    first_analysis = response.text

                                    # 문제점 키워드 탐지
                                    problem_keywords = [
                                        "위반", "문제", "충돌", "부적절", "개선", "수정", "보완",
                                        "법령 위반", "상위법령", "위법", "불일치", "모순", "우려"
                                    ]

                                    has_problems = any(keyword in first_analysis for keyword in problem_keywords)

                                    if has_problems:
                                        st.warning(f"⚠️ Gemini가 잠재적 문제점을 발견했습니다!")

                                    # 🆕 3-2단계: Gemini 분석 결과 기반 정밀 재검색
                                    # ✅ 위법성 유무와 관계없이 항상 검색 (유사 사례도 참고 가치 있음)
                                    if st.session_state.gemini_store_manager:
                                        with st.spinner("🔍 1차 분석 결과를 기반으로 관련 판례를 정밀 검색하고 있습니다..."):
                                            try:
                                                # 핵심 키워드 추출 (조례명 + 구체적인 조항 제목)
                                                # re 모듈은 파일 상단에서 이미 import됨

                                                # 1. 조례명이 있으면 사용
                                                search_keywords = []
                                                if ordinance_name:
                                                    search_keywords.append(ordinance_name)

                                                # 2. 분석 결과에서 제○조 패턴 추출
                                                article_mentions = re.findall(r'제\s*\d+\s*조[^,\n]{0,30}', first_analysis)
                                                search_keywords.extend(article_mentions[:5])

                                                # 3. 핵심 법적 쟁점 키워드 추출
                                                key_issues = []
                                                issue_patterns = [
                                                    r'(기관위임사무)',
                                                    r'(직업선택의\s*자유)',
                                                    r'(계약[의]?\s*자유)',
                                                    r'(법률유보[원칙]?)',
                                                    r'(평등권)',
                                                    r'(재산권)',
                                                    r'(영업의\s*자유)',
                                                    r'(과잉금지[원칙]?)',
                                                ]
                                                for pattern in issue_patterns:
                                                    matches = re.findall(pattern, first_analysis)
                                                    key_issues.extend(matches)

                                                # 중복 제거
                                                key_issues = list(set(key_issues))[:5]

                                                # 1️⃣ 판례 및 사례 검색 쿼리 생성 (간결하게)
                                                if ordinance_name and key_issues:
                                                    # 조례명 + 법적 쟁점
                                                    case_query = f"'{ordinance_name}'과 관련된 {', '.join(key_issues)} 위반 판례와 재의·제소 사례를 찾아주세요."
                                                elif ordinance_name:
                                                    # 조례명만
                                                    case_query = f"'{ordinance_name}'의 위법 판례, 재의 요구, 제소 사례를 찾아주세요."
                                                elif key_issues:
                                                    # 법적 쟁점만
                                                    case_query = f"{', '.join(key_issues)} 위반 조례 판례와 재의·제소 사례를 찾아주세요."
                                                else:
                                                    # 일반 검색
                                                    case_query = "조례 위법 판례와 재의·제소 사례를 찾아주세요."

                                                # 2️⃣ 이론적 설명 및 가이드라인 검색 쿼리 생성
                                                if key_issues:
                                                    # 구체적인 법적 쟁점이 있는 경우
                                                    theory_query = f"{', '.join(key_issues)}에 대한 법리, 이론적 설명, 판단 기준을 설명해주세요."
                                                else:
                                                    # 일반적인 조례 제정 이론 검색
                                                    theory_query = "조례 제정의 법리와 원칙, 상위법령 위배 판단 기준을 설명해주세요."

                                                # 판례/사례 검색 수행
                                                case_result = st.session_state.gemini_store_manager.search(
                                                    case_query,
                                                    top_k=5
                                                )

                                                # 이론/가이드라인 검색 수행
                                                theory_result = st.session_state.gemini_store_manager.search(
                                                    theory_query,
                                                    top_k=5
                                                )

                                                # 검색 결과 통합
                                                case_answer = case_result.get('answer', '')
                                                case_sources = case_result.get('sources', [])

                                                theory_answer = theory_result.get('answer', '')
                                                theory_sources = theory_result.get('sources', [])

                                                # 두 검색 결과를 결합
                                                combined_answer = ""
                                                combined_sources = []

                                                if case_answer and len(case_answer) > 200:
                                                    combined_answer += "## 📚 관련 판례 및 재의·제소 사례\n\n"
                                                    combined_answer += case_answer
                                                    combined_sources.extend(case_sources)

                                                if theory_answer and len(theory_answer) > 200:
                                                    if combined_answer:
                                                        combined_answer += "\n\n---\n\n"
                                                    combined_answer += "## 📖 이론적 근거 및 법리 해설\n\n"
                                                    combined_answer += theory_answer
                                                    combined_sources.extend(theory_sources)

                                                # 최종 답변 설정
                                                refined_answer = combined_answer if combined_answer else ""
                                                refined_sources = combined_sources

                                                if refined_answer and len(refined_answer) > 500:
                                                    # 기존 판례 결과에 추가
                                                    search_summary = []
                                                    if case_answer and len(case_answer) > 200:
                                                        search_summary.append(f"판례/사례 {len(case_answer)}자")
                                                    if theory_answer and len(theory_answer) > 200:
                                                        search_summary.append(f"이론/법리 {len(theory_answer)}자")

                                                    refined_case = {
                                                        'violation_type': '정밀 검색 결과 (판례 + 이론)',
                                                        'content': refined_answer,
                                                        'similarity': 0.98,
                                                        'topic': f'정밀 검색: 판례·사례 및 이론적 근거 ({", ".join(search_summary)})',
                                                        'relevance_score': 0.98,
                                                        'context_relevance': 0.95,
                                                        'matched_concepts': ['판례', '이론', '법리', '가이드라인', '정밀검색'],
                                                        'summary': refined_answer[:200] + '...',
                                                        'metadata': {
                                                            'source': 'gemini_file_search_comprehensive',
                                                            'source_files': [s.get('title', '') for s in refined_sources if s.get('title')],
                                                            'query_case': case_query,
                                                            'query_theory': theory_query,
                                                            'search_type': 'comprehensive_analysis_based',
                                                            'has_cases': bool(case_answer and len(case_answer) > 200),
                                                            'has_theory': bool(theory_answer and len(theory_answer) > 200)
                                                        }
                                                    }

                                                    # 정밀 검색 결과를 맨 앞에 추가 (가장 관련성 높음)
                                                    theoretical_results.insert(0, refined_case)
                                                    st.session_state.theoretical_results = theoretical_results

                                                    st.success(f"✅ 분석 결과 기반 정밀 검색 완료: {', '.join(search_summary)}")

                                                    # 미리보기
                                                    with st.expander("🎯 정밀 검색 결과 미리보기 (판례 + 이론)", expanded=True):
                                                        st.markdown(f"**{refined_case['topic']}**")
                                                        st.markdown(f"📄 {refined_answer[:500]}...")

                                                        # 출처 파일 표시
                                                        unique_sources = list(set([s for s in refined_case['metadata']['source_files'] if s]))
                                                        if unique_sources:
                                                            st.markdown(f"📁 출처: {', '.join(unique_sources[:5])}")

                                                        # 검색 유형 표시
                                                        if refined_case['metadata']['has_cases']:
                                                            st.markdown("✓ 판례 및 재의·제소 사례 포함")
                                                        if refined_case['metadata']['has_theory']:
                                                            st.markdown("✓ 이론적 근거 및 법리 해설 포함")
                                                else:
                                                    st.info("ℹ️ 정밀 검색에서 추가 판례를 찾지 못했습니다.")

                                            except Exception as e:
                                                st.warning(f"⚠️ 정밀 검색 중 오류 (계속 진행): {str(e)}")

                                    else:
                                        st.success("✅ Gemini 1차 분석에서 특별한 문제점이 발견되지 않았습니다.")

                                    analysis_results.append({
                                        'model': 'Gemini (1차 분석)',
                                        'content': first_analysis
                                    })
                                else:
                                    st.error("Gemini 1차 분석 응답이 비어있습니다.")
                            except Exception as e:
                                st.error(f"Gemini 1차 분석 오류: {str(e)}")
                                analysis_results.append({
                                    'model': 'Gemini (1차 분석)',
                                    'error': str(e)
                                })
                        
                        # 4단계: 문제 발견 시 자료 참고 분석 수행
                        relevant_guidelines = None
                        loaded_stores = []
                        enhanced_analysis = None
                        
                        if has_problems and use_auto_search and first_analysis:
                            # 4단계: Gemini File Search를 사용한 관련 자료 검색
                            comprehensive_analysis_results = None

                            # 발견된 문제점을 기반으로 구체적인 검색 쿼리 생성
                            search_terms = []

                            # 사무 관련 문제
                            if any(word in first_analysis for word in ["소관사무", "사무구분", "위임사무", "자치사무"]):
                                search_terms.extend(["기관위임사무 조례제정 불가", "위임사무 조례 제정 한계"])

                            # 법령 위반 관련 문제
                            if any(word in first_analysis for word in ["법령 위반", "상위법령", "법령우위", "위반"]):
                                search_terms.extend(["법령 위반 조례 사례", "상위법령 충돌 조례"])

                            # 조례 제정 한계 관련
                            if any(word in first_analysis for word in ["제정 한계", "입법한계", "불가", "위법"]):
                                search_terms.extend(["조례 제정 한계 판례", "위법 조례 제정 사례"])

                            # 기본 검색어가 없으면 일반적인 검색어 사용
                            if not search_terms:
                                search_terms = ["법령 위반 조례 판례", "조례 제정 한계 사례"]

                            # 여러 검색어 중 하나 선택 (가장 구체적인 것)
                            search_query = search_terms[0] if search_terms else "위법 조례 판례"

                            # Gemini File Search 사용
                            if st.session_state.gemini_store_manager:
                                try:
                                    relevant_guidelines = search_relevant_guidelines_gemini(
                                        query=search_query,
                                        api_key=gemini_api_key,
                                        store_manager=st.session_state.gemini_store_manager,
                                        top_k=8
                                    )
                                    loaded_stores = ["Gemini File Search (통합 저장소)"]

                                    if relevant_guidelines:
                                        st.success(f"✅ {len(relevant_guidelines)}개의 관련 자료를 발견했습니다")

                                except Exception as e:
                                    st.error(f"Gemini 검색 오류: {e}")
                                    relevant_guidelines = []
                                    loaded_stores = []
                            else:
                                st.warning("⚠️ Gemini File Search가 초기화되지 않았습니다. API 키를 확인해주세요.")
                                relevant_guidelines = []
                                loaded_stores = []
                            
                            if relevant_guidelines and loaded_stores:
                                st.success(f"✅ {len(loaded_stores)}개 자료에서 {len(relevant_guidelines)}개 관련 내용을 검색했습니다:")
                                for store in loaded_stores:
                                    st.markdown(f"   • {store}")
                                
                                # 가이드라인 미리보기 (선택사항)
                                with st.expander("📖 검색된 문제 관련 자료 미리보기", expanded=False):
                                    source_groups = {}
                                    for guideline in relevant_guidelines:
                                        source_store = guideline.get('source_store', '알 수 없는 자료')
                                        if source_store not in source_groups:
                                            source_groups[source_store] = []
                                        source_groups[source_store].append(guideline)
                                    
                                    for source_store, guidelines in source_groups.items():
                                        st.markdown(f"**📚 {source_store}**")
                                        for i, guideline in enumerate(guidelines):
                                            similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                                            st.markdown(f"   [{i+1}] (유사도: {similarity_score:.3f})")
                                            st.markdown(guideline['text'][:200] + "..." if len(guideline['text']) > 200 else guideline['text'])
                                            st.markdown("---")
                                
                                # 2차 보강 분석 수행 (조용히) - Ollama Cloud 우선 사용
                                if has_ollama:
                                    try:
                                        # 보강 분석용 프롬프트
                                        enhanced_prompt = create_analysis_prompt(
                                            pdf_text,
                                            search_results_for_analysis,
                                            superior_laws_content,
                                            relevant_guidelines,
                                            is_first_ordinance,
                                            comprehensive_analysis_results,
                                            theoretical_results
                                        )

                                        with st.spinner("🤖 AI가 참고 자료를 바탕으로 보강 분석을 수행하고 있습니다..."):
                                            enhanced_analysis = call_ollama_cloud_api(enhanced_prompt)

                                        if enhanced_analysis:
                                            analysis_results.append({
                                                'model': f'Ollama Cloud (자료 참고 보강분석 - {len(loaded_stores)}개 자료)',
                                                'content': enhanced_analysis
                                            })
                                    except Exception as e:
                                        st.error(f"Ollama Cloud 보강 분석 오류: {str(e)}")
                                elif gemini_api_key:
                                    try:
                                        # 보강 분석용 프롬프트
                                        enhanced_prompt = create_analysis_prompt(
                                            pdf_text,
                                            search_results_for_analysis,
                                            superior_laws_content,
                                            relevant_guidelines,
                                            is_first_ordinance,
                                            comprehensive_analysis_results,
                                            theoretical_results
                                        )

                                        enhanced_response = model.generate_content(enhanced_prompt)
                                        if enhanced_response and hasattr(enhanced_response, 'text') and enhanced_response.text:
                                            enhanced_analysis = enhanced_response.text
                                            analysis_results.append({
                                                'model': f'Gemini (자료 참고 보강분석 - {len(loaded_stores)}개 자료)',
                                                'content': enhanced_analysis
                                            })
                                    except Exception as e:
                                        st.error(f"자료 참고 보강 분석 오류: {str(e)}")
                            else:
                                st.info("문제점과 관련된 자료를 찾지 못했습니다.")
                        elif not has_problems:
                            st.info("✅ 문제점이 발견되지 않아 자료 검색을 건너뜁니다.")
                        elif not use_auto_search:
                            st.info("🔄 자동 참고 자료 검색 기능이 비활성화되어 있습니다.")
                        
                        # 5단계: OpenAI 추가 분석 (선택사항)
                        if openai_api_key:
                            try:
                                openai.api_key = openai_api_key
                                # 가장 완전한 프롬프트로 OpenAI 분석
                                openai_prompt = create_analysis_prompt(pdf_text, search_results_for_analysis, superior_laws_content, relevant_guidelines, is_first_ordinance, comprehensive_analysis_results, theoretical_results)
                                
                                response = openai.ChatCompletion.create(
                                    model="gpt-4o-mini",
                                    messages=[
                                        {"role": "system", "content": "당신은 법률 전문가입니다. 조례 분석과 검토를 도와주세요."},
                                        {"role": "user", "content": openai_prompt}
                                    ],
                                    temperature=0.7,
                                    max_tokens=4000
                                )
                                
                                if response.choices[0].message.content:
                                    analysis_results.append({
                                        'model': 'OpenAI (추가 분석)',
                                        'content': response.choices[0].message.content
                                    })
                            except Exception as e:
                                st.error(f"OpenAI 분석 오류: {str(e)}")
                                analysis_results.append({
                                    'model': 'OpenAI (추가 분석)',
                                    'error': str(e)
                                })
                        
                        if analysis_results:
                            # 🆕 분석 결과를 세션 상태에 저장
                            st.session_state.analysis_results = analysis_results
                            st.session_state.analysis_metadata = {
                                'has_problems': has_problems,
                                'relevant_guidelines': relevant_guidelines,
                                'loaded_stores': loaded_stores,
                                'is_first_ordinance': is_first_ordinance,
                                'superior_laws_content': superior_laws_content,
                                'search_results_for_analysis': search_results_for_analysis,
                                'pdf_text': pdf_text,
                                'analysis_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            }

                            # 분석 완료 메시지
                            st.markdown("---")
                            if has_problems and relevant_guidelines and loaded_stores:
                                st.success(f"🎯 **복합 자료 보강 분석 완료**: 문제점 탐지 → {len(loaded_stores)}개 자료 참고 → 보강 분석")
                            elif has_problems and relevant_guidelines:
                                st.success("🎯 **지능형 분석 완료**: 문제점 탐지 → 자료 검색 → 보강 분석")
                            elif has_problems:
                                st.info("⚠️ **문제점 탐지 분석 완료**: 자료 검색 없이 기본 분석만 수행")
                            else:
                                st.success("✅ **기본 분석 완료**: 특별한 문제점이 발견되지 않음")
                            
                            # 분석 결과 요약
                            analysis_count = len([r for r in analysis_results if 'error' not in r])
                            error_count = len([r for r in analysis_results if 'error' in r])
                            
                            if analysis_count > 0:
                                # 🆕 선택된 조례 수 정확히 반영
                                if is_first_ordinance:
                                    analysis_type_text = "최초 제정 조례"
                                else:
                                    selected_count = len(search_results_for_analysis)
                                    analysis_type_text = f"선택된 {selected_count}개 타 시도 조례 비교"
                                st.markdown(f"**📋 분석 유형**: {analysis_type_text}")
                                st.markdown(f"**🤖 수행된 분석**: {analysis_count}개")
                                if relevant_guidelines:
                                    st.markdown(f"**📚 참고된 가이드라인**: {len(relevant_guidelines)}개")
                            
                            # 최종 보고서만 표시 (자료 참고 보강 분석 또는 OpenAI 분석)
                            final_report = None

                            # 우선순위: 자료 참고 보강분석 > OpenAI 추가 분석 > 1차 분석
                            for result in reversed(analysis_results):  # 역순으로 최신 결과 우선
                                if 'error' not in result:
                                    if "자료 참고 보강분석" in result['model']:
                                        final_report = result
                                        break
                                    elif "자료 참고" in result['model'] or "OpenAI" in result['model']:
                                        final_report = result
                                        break

                            # 자료 참고나 OpenAI가 없으면 1차 분석 사용
                            if not final_report:
                                for result in analysis_results:
                                    if 'error' not in result and "1차 분석" in result['model']:
                                        final_report = result
                                        break

                            # 최종 보고서 표시
                            if final_report:
                                st.markdown("### 📋 최종 분석 보고서")

                                # 보고서 타입 표시
                                if "자료 참고 보강분석" in final_report['model']:
                                    st.success("🎯 **자료 참고 보강 분석 결과**")
                                    st.caption(f"📚 **활용 모델**: {final_report['model']}")
                                elif "자료 참고" in final_report['model']:
                                    st.success("🎯 **참고 자료 기반 보강 분석 결과**")
                                elif "OpenAI" in final_report['model']:
                                    st.info("📊 **OpenAI 추가 분석 결과**")
                                elif "Ollama Cloud" in final_report['model']:
                                    st.info("🤖 **Ollama Cloud AI 분석 결과** (무료 서비스)")
                                else:
                                    st.info("🤖 **Gemini 기본 분석 결과**")

                                # 보고서 내용 (content 또는 analysis 키 지원)
                                report_content = final_report.get('content') or final_report.get('analysis', '')
                                st.markdown(report_content)

                            # 오류 메시지만 별도 표시
                            for result in analysis_results:
                                if 'error' in result:
                                    st.error(f"❌ {result['model']} 오류: {result['error']}")
                            
                            # Word 문서 생성 및 다운로드
                            with st.spinner("분석 결과 Word 문서 생성 중..."):
                                doc = create_comparison_document(pdf_text, search_results_for_analysis, analysis_results, superior_laws_content, relevant_guidelines)
                                
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_bytes = doc_io.getvalue()
                                
                                # 파일명에 분석 방식 표시
                                if has_problems and relevant_guidelines and loaded_stores:
                                    stores_count = len(loaded_stores)
                                    filename_prefix = f"복합자료보강분석({stores_count}개자료)" if is_first_ordinance else f"조례비교_복합자료분석({stores_count}개자료)"
                                elif has_problems and relevant_guidelines:
                                    filename_prefix = "자료참고보강분석" if is_first_ordinance else "조례비교_자료분석"
                                elif has_problems:
                                    filename_prefix = "문제점탐지분석" if is_first_ordinance else "조례비교_문제점분석"
                                else:
                                    filename_prefix = "최초조례_기본분석" if is_first_ordinance else "조례_기본비교분석"
                                
                                st.download_button(
                                    label="📄 분석 결과 Word 문서 다운로드",
                                    data=doc_bytes,
                                    file_name=f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        else:
                            st.error("분석 결과가 없습니다.")

    # --------------------------------------------------------------------------
    # 오류 제보 기능 (AI 분석 버튼과 분리하여 하단에 배치)
    # --------------------------------------------------------------------------
    st.divider()
    with st.expander("🚨 AI 분석 오류 제보 (Report Error)", expanded=False):
        st.markdown("AI 분석 결과가 이상하거나 오류가 발생한 경우 제보해주세요.")
        
        with st.form("error_report_form"):
            report_content = st.text_area("오류 내용 / 불편 사항", height=150, placeholder="구체적인 오류 내용이나 개선 요청사항을 적어주세요.")
            uploaded_file = st.file_uploader("화면 캡처 첨부 (선택사항)", type=['png', 'jpg', 'jpeg'])
            
            submit_report = st.form_submit_button("제보하기")
            
            if submit_report:
                if not report_content:
                    st.warning("내용을 입력해주세요.")
                else:
                    with st.spinner("제보 내용을 전송하고 있습니다..."):
                        # 첨부파일 처리
                        attachment_data = None
                        attachment_name = None
                        if uploaded_file is not None:
                            attachment_data = uploaded_file.getvalue()
                            attachment_name = uploaded_file.name
                        
                        # 이메일 전송
                        subject = f"[조례분석AI] 오류 제보: {report_content[:20]}..."
                        body = f"내용:\n{report_content}\n\n(첨부파일 있음)" if attachment_data else f"내용:\n{report_content}"
                        
                        success = send_error_report(subject, body, attachment_data, attachment_name)
                        
                        if success:
                            st.success("✅ 제보가 성공적으로 전송되었습니다. 소중한 의견 감사합니다!")
                        else:
                            st.error("❌ 전송에 실패했습니다. 잠시 후 다시 시도해주세요.")

if __name__ == "__main__":
    main()